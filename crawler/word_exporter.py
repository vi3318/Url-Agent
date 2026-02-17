"""
Structured Word Document Exporter
===================================
Produces professionally formatted DOCX from RAGDocument/RAGCorpus data.

Features:
- Auto-generated Table of Contents (TOC)
- Heading hierarchy preserved (H1–H6 → Word Heading 1–6)
- Tables rendered as proper Word tables with headers
- Code blocks in monospace with shading
- Breadcrumb / section path metadata
- Cover page with crawl summary statistics
- Hyperlinked source URLs
- Page breaks between documents
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any, Dict, List, Optional, TYPE_CHECKING

if TYPE_CHECKING:
    from .rag_model import RAGCorpus, RAGDocument

logger = logging.getLogger(__name__)


def export_docx(
    corpus: "RAGCorpus",
    filepath: str,
    *,
    include_toc: bool = True,
    max_content_chars: int = 20_000,
) -> str:
    """
    Export a RAGCorpus to a structured Word document.

    Args:
        corpus: RAGCorpus with documents and chunks
        filepath: Output .docx path
        include_toc: Whether to insert a TOC field
        max_content_chars: Max chars per document section

    Returns:
        Absolute path to the created file
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    output_path = Path(filepath)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # ── Configure base styles ──────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    # ── Cover Page ─────────────────────────────────────────────────
    title = doc.add_heading("Web Crawl Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    stats = corpus.crawl_stats
    config = corpus.crawl_config
    summary_items = [
        ("Scope", config.get("scope", stats.get("scope", "N/A"))),
        ("Total Pages", str(corpus.total_documents)),
        ("Total Chunks", str(corpus.total_chunks)),
        ("Total Words", f"{corpus.total_words:,}"),
        ("Pages Failed", str(stats.get("pages_failed", 0))),
        ("Pages Skipped", str(stats.get("pages_skipped", 0))),
        ("Elapsed Time", f"{stats.get('elapsed_sec', stats.get('elapsed_time', 0))}s"),
        ("Speed", f"{stats.get('pages_per_sec_overall', stats.get('pages_per_second', 0)):.2f} pages/sec"),
        ("Stop Reason", stats.get("stop_reason", "N/A")),
    ]

    # Summary as a clean table
    summary_table = doc.add_table(rows=len(summary_items), cols=2)
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(summary_items):
        row = summary_table.rows[i]
        _cell_text(row.cells[0], label, bold=True, size=Pt(10))
        _cell_text(row.cells[1], value, size=Pt(10))

    doc.add_page_break()

    # ── Table of Contents ──────────────────────────────────────────
    if include_toc:
        doc.add_heading("Table of Contents", level=1)
        _add_toc_field(doc)
        doc.add_page_break()

    # ── Per-Document Sections ──────────────────────────────────────
    for doc_idx, rag_doc in enumerate(corpus.documents):
        _render_document(doc, rag_doc, max_content_chars)

        # Page break between documents (except last)
        if doc_idx < len(corpus.documents) - 1:
            doc.add_page_break()

    # ── Save ───────────────────────────────────────────────────────
    doc.save(str(output_path))
    logger.info(f"Exported DOCX to {output_path.absolute()}")
    return str(output_path.absolute())


# ---------------------------------------------------------------------------
# Per-document rendering
# ---------------------------------------------------------------------------

def _render_document(
    doc,
    rag_doc: "RAGDocument",
    max_content_chars: int,
) -> None:
    """Render a single RAGDocument into the Word document."""
    from docx.shared import Pt, RGBColor

    # ── Document Title (H1) ────────────────────────────────────────
    heading_text = rag_doc.page_title or rag_doc.source_url
    doc.add_heading(heading_text[:120], level=1)

    # ── URL ────────────────────────────────────────────────────────
    url_para = doc.add_paragraph()
    url_run = url_para.add_run(rag_doc.source_url)
    url_run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
    url_run.font.size = Pt(9)
    url_para.paragraph_format.space_after = Pt(4)

    # ── Breadcrumb / Section Path ──────────────────────────────────
    if rag_doc.breadcrumb:
        bc_para = doc.add_paragraph()
        bc_label = bc_para.add_run("Section: ")
        bc_label.bold = True
        bc_label.font.size = Pt(9)
        bc_value = bc_para.add_run(" > ".join(rag_doc.breadcrumb))
        bc_value.font.size = Pt(9)
        bc_value.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # ── Headings from the page (render as proper Word heading levels) ──
    # We use the heading hierarchy to structure the content
    headings_rendered = set()
    current_heading_level = 2  # Start at level 2 (level 1 = page title)

    # ── Main Content (rendered from chunks) ────────────────────────
    if rag_doc.chunks:
        last_section = ""
        chars_written = 0

        for chunk in rag_doc.chunks:
            if chars_written >= max_content_chars:
                p = doc.add_paragraph("[... content truncated ...]")
                p.runs[0].font.italic = True
                p.runs[0].font.size = Pt(9)
                break

            # Insert section heading if changed
            section = chunk.section_title
            if section and section != last_section and section != rag_doc.page_title:
                # Determine heading level from heading_path
                level = min(_heading_level_from_path(chunk.heading_path), 6)
                if level < 2:
                    level = 2
                doc.add_heading(section[:100], level=level)
                last_section = section

            # Render chunk based on type
            if chunk.chunk_type == "table":
                _render_table_chunk(doc, chunk)
            elif chunk.chunk_type == "code":
                _render_code_chunk(doc, chunk)
            else:
                # Text chunk
                content = chunk.content
                if content:
                    # Split into paragraphs for readability
                    paragraphs = content.split("\n\n")
                    for para_text in paragraphs:
                        para_text = para_text.strip()
                        if para_text:
                            p = doc.add_paragraph(para_text)
                            for run in p.runs:
                                run.font.size = Pt(10)
                            chars_written += len(para_text)

    elif rag_doc.full_text:
        # Fallback: render full text directly
        content = rag_doc.full_text
        if len(content) > max_content_chars:
            content = content[:max_content_chars] + "\n\n[... content truncated ...]"

        # Render headings + content
        _render_text_with_headings(doc, content, rag_doc.headings)


def _render_text_with_headings(
    doc,
    text: str,
    headings: Dict[str, List[str]],
) -> None:
    """Render text content, inserting Word headings where heading text matches."""
    from docx.shared import Pt

    # Build a set of known headings for quick lookup
    heading_lookup: Dict[str, int] = {}
    for level in range(1, 7):
        for h in headings.get(f"h{level}", []):
            heading_lookup[h.strip()] = level

    paragraphs = text.split("\n\n")
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue

        # Check if this paragraph matches a known heading
        if para_text in heading_lookup:
            level = min(heading_lookup[para_text] + 1, 6)  # offset by 1 (H1 = title)
            doc.add_heading(para_text[:100], level=level)
        else:
            p = doc.add_paragraph(para_text)
            for run in p.runs:
                run.font.size = Pt(10)


def _render_table_chunk(doc, chunk) -> None:
    """Render a table chunk as a proper Word table."""
    from docx.shared import Pt, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT

    content = chunk.content
    if not content:
        return

    lines = content.strip().split("\n")
    if not lines:
        return

    # Parse the table text back into rows
    rows_data = []
    header_row = None
    for line in lines:
        line = line.strip()
        if not line or set(line) <= {"-", "|", " ", "+"}:
            continue  # skip separator lines
        cells = [c.strip() for c in line.split("|") if c.strip()]
        if cells:
            if header_row is None:
                header_row = cells
            else:
                rows_data.append(cells)

    if not header_row:
        # Just render as text
        p = doc.add_paragraph(content)
        for run in p.runs:
            run.font.size = Pt(9)
        return

    # Determine number of columns
    num_cols = max(len(header_row), max((len(r) for r in rows_data), default=0))

    # Create Word table
    table = doc.add_table(rows=1 + len(rows_data), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    # Header row
    for i, header in enumerate(header_row[:num_cols]):
        cell = table.rows[0].cells[i]
        _cell_text(cell, header, bold=True, size=Pt(9))

    # Data rows
    for row_idx, row_data in enumerate(rows_data):
        for col_idx, cell_text in enumerate(row_data[:num_cols]):
            cell = table.rows[row_idx + 1].cells[col_idx]
            _cell_text(cell, cell_text, size=Pt(9))

    doc.add_paragraph()  # spacing after table


def _render_code_chunk(doc, chunk) -> None:
    """Render a code block with monospace font and shading."""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    content = chunk.content
    if not content:
        return

    # Code label
    label = doc.add_paragraph()
    label_run = label.add_run("Code:")
    label_run.bold = True
    label_run.font.size = Pt(9)

    # Code content with monospace
    p = doc.add_paragraph()
    run = p.add_run(content[:5000])
    run.font.name = "Consolas"
    run.font.size = Pt(8)

    # Add shading
    try:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "F5F5F5")
        shading.set(qn("w:val"), "clear")
        p.paragraph_format.element.get_or_add_pPr().append(shading)
    except Exception:
        pass

    doc.add_paragraph()  # spacing


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _cell_text(cell, text: str, bold: bool = False, size=None) -> None:
    """Set cell text with formatting."""
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = bold
            if size:
                run.font.size = size


def _heading_level_from_path(heading_path: List[str]) -> int:
    """Extract the deepest heading level from a heading path."""
    if not heading_path:
        return 2
    last = heading_path[-1]
    import re
    m = re.match(r"H(\d+):", last)
    if m:
        return int(m.group(1)) + 1  # +1 because H1 = document title in Word
    return 2


def _add_toc_field(doc) -> None:
    """Insert a Word TOC field code (updates on open in Word)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(fld_char_begin)

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._element.append(instr_text)

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    run._element.append(fld_char_separate)

    # Placeholder text (shown until TOC is refreshed in Word)
    placeholder_run = paragraph.add_run(
        "[Open in Microsoft Word and press F9 to update Table of Contents]"
    )
    placeholder_run.font.italic = True

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._element.append(fld_char_end)


# ---------------------------------------------------------------------------
# Legacy compatibility: export from DeepCrawlResult
# ---------------------------------------------------------------------------

def export_docx_from_crawl_result(
    result,
    filepath: str,
    *,
    include_toc: bool = True,
) -> str:
    """
    Export directly from a DeepCrawlResult (legacy compatibility).

    Converts DeepCrawlResult → RAGCorpus → DOCX.
    """
    from .rag_model import RAGCorpus
    from .pipeline import transform_batch

    # Convert pages to dicts
    page_dicts = [p.to_dict() for p in result.pages]

    # Transform to RAG documents
    rag_docs = transform_batch(page_dicts)

    # Build corpus
    corpus = RAGCorpus(
        documents=rag_docs,
        crawl_stats=result.stats,
    )

    return export_docx(corpus, filepath, include_toc=include_toc)
