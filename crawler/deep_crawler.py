"""
Deep Interactive Crawler
Specialized crawler for hierarchical documentation sites with expandable content.
Clicks dropdowns, expands accordions, and crawls revealed sub-links.

SAFETY FEATURES:
- Strict per-page timeout (20s default)
- Per-page click limit (50 clicks max)
- Static HTML fallback when JS rendering fails
- Graceful shutdown when limits are reached
- No networkidle waits (uses domcontentloaded)
"""

import logging
import time
import json
import re
import requests
from typing import List, Dict, Set, Optional, Callable, Any, Tuple
from dataclasses import dataclass, field
from urllib.parse import urlparse, urljoin
from pathlib import Path
from bs4 import BeautifulSoup

from playwright.sync_api import sync_playwright, Page, Browser, BrowserContext
from playwright.sync_api import TimeoutError as PlaywrightTimeout

from .scraper import PageData, PageScraper
from .scope_filter import ScopeFilter
from .utils import URLNormalizer, ensure_joinable_base
from . import interaction_policy

logger = logging.getLogger(__name__)

# Best-available HTML parser for BeautifulSoup (static fallback path)
try:
    import lxml  # noqa: F401
    _BS_PARSER = "lxml"
except ImportError:
    _BS_PARSER = "html.parser"

@dataclass
class DeepCrawlConfig:
    """
    Configuration for deep interactive crawling.
    
    SAFETY LIMITS (enforced strictly to prevent hangs):
    - Per-page timeout: 20 seconds max
    - Wait strategy: domcontentloaded (NOT networkidle)
    - Max expandable clicks per page: 50
    - Max pages: 150 (default)
    - Max depth: 5 (default)
    """
    # Crawl limits - STRICT DEFAULTS
    max_pages: int = 150  # Reduced for stability
    max_depth: int = 5    # Reasonable depth limit
    timeout: int = 20000  # 20 seconds - strict per-page timeout (ms)
    
    # Rate limiting
    delay_between_pages: float = 1.0  # seconds
    delay_after_click: float = 0.3   # reduced for efficiency
    
    # Per-page expansion limits - CRITICAL for preventing hangs
    max_clicks_per_page: int = 50  # Stop expansion after 50 clicks per page
    
    # Browser settings
    headless: bool = True
    viewport_width: int = 1920
    viewport_height: int = 1080
    
    # User agent
    user_agent: str = (
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
        "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36"
    )
    
    # Enable static fallback when JS rendering fails
    enable_static_fallback: bool = True
    
    # GENERALIZED interactive element selectors
    # Default: None → uses interaction_policy.DEFAULT_INTERACTIVE_SELECTORS
    # which covers 20+ frameworks (Bootstrap, React, Docusaurus, MkDocs, etc.)
    interactive_selectors: Optional[List[str]] = None
    
    # Link selectors to find after expansion
    link_selectors: List[str] = field(default_factory=lambda: [
        'a[href]',
        # Tree / sidebar navigation
        '.toc-link[href]',
        '.nav-link[href]',
        '[role="treeitem"] a',
        # Docusaurus
        '.menu__link[href]',
        # ReadTheDocs / Sphinx
        '.toctree-l1 a[href]',
        '.toctree-l2 a[href]',
        # MkDocs
        '.md-nav__link[href]',
        # Confluence
        '.expand-content a[href]',
        # GitBook
        'nav a[href]',
    ])
    
    # Content selectors (main content area)
    content_selectors: List[str] = field(default_factory=lambda: [
        'main',
        'article', 
        '.content',
        '.main-content',
        '#content',
        '[role="main"]',
        # Documentation platforms
        '.documentation',
        '.doc-content',
        '.ohc-main-content',          # Oracle
        '.topic-content',             # Oracle
        '.theme-doc-markdown',        # Docusaurus
        '.markdown-section',          # Docusaurus
        '.md-content',                # MkDocs Material
        '.rst-content',               # ReadTheDocs
        '.document',                  # Sphinx
        '.wiki-content',              # Confluence
        '.article-body',              # Zendesk Guide
        # React UI library wrappers
        '.MuiContainer-root',         # MUI
        '.chakra-container',          # Chakra
        # Enterprise
        '.slds-template__container',  # Salesforce
    ])
    
    # Selectors to exclude from content
    exclude_selectors: List[str] = field(default_factory=lambda: [
        'nav',
        'header',
        'footer',
        '.sidebar',
        '.toc',
        '.breadcrumb',
        'script',
        'style',
        'noscript',
    ])

    # Scope-filter params (populated by CrawlerRunConfig converters)
    deny_patterns: List[str] = field(default_factory=list)
    strip_all_queries: bool = False

    # Built-in URL deny patterns — always applied to filter junk URLs
    # These catch image viewers, attachments, non-English locale duplicates, etc.
    builtin_deny_patterns: List[str] = field(default_factory=lambda: [
        r'/viewer/attachment/',       # image/file viewer pages
        r'/viewer/',                  # generic viewer pages
        r'/(de-DE|fr-FR|ko-KR|ja-JP|zh-CN|zh-TW|pt-BR|es-ES|it-IT|nl-NL|ru-RU|pl-PL|sv-SE|da-DK|fi-FI|nb-NO|cs-CZ|hu-HU|ro-RO|tr-TR|th-TH|he-IL|ar-SA|id-ID|ms-MY|vi-VN|uk-UA|el-GR|bg-BG|hr-HR|sk-SK|sl-SI|lt-LT|lv-LV|et-EE)/',  # non-English locale variants
    ])

    # Minimum word count to accept a page into results
    # Pages below this threshold are considered empty/junk
    min_word_count: int = 10


@dataclass  
class DeepPageData:
    """Data extracted from a deeply crawled page."""
    url: str
    title: str = ""
    breadcrumb: List[str] = field(default_factory=list)
    headings: Dict[str, List[str]] = field(default_factory=dict)
    text_content: str = ""
    tables: List[Dict] = field(default_factory=list)
    code_blocks: List[str] = field(default_factory=list)
    internal_links: List[str] = field(default_factory=list)
    parent_url: str = ""
    depth: int = 0
    section_path: List[str] = field(default_factory=list)  # Hierarchy path like ["3 Absence Management", "Tables", "ANC_ABSENCE_AGREEMENTS_F"]
    word_count: int = 0
    _skipped: bool = False  # True if page was empty/cookie-only (excluded from results)
    
    def to_dict(self) -> dict:
        return {
            'url': self.url,
            'title': self.title,
            'breadcrumb': self.breadcrumb,
            'section_path': self.section_path,
            'headings': self.headings,
            'text_content': self.text_content,
            'tables': self.tables,
            'code_blocks': self.code_blocks,
            'internal_links': self.internal_links,
            'parent_url': self.parent_url,
            'depth': self.depth,
            'word_count': self.word_count,
        }
    
    def to_flat_dict(self) -> dict:
        """Flat dict for CSV export."""
        return {
            'url': self.url,
            'title': self.title,
            'breadcrumb': ' > '.join(self.breadcrumb),
            'section_path': ' > '.join(self.section_path),
            'h1': ' | '.join(self.headings.get('h1', [])),
            'h2': ' | '.join(self.headings.get('h2', [])),
            'h3': ' | '.join(self.headings.get('h3', [])),
            'text_content': self.text_content[:15000],
            'tables_count': len(self.tables),
            'code_blocks_count': len(self.code_blocks),
            'internal_links_count': len(self.internal_links),
            'depth': self.depth,
            'word_count': self.word_count,
        }


@dataclass
class DeepCrawlResult:
    """Result of a deep crawl operation."""
    pages: List[DeepPageData] = field(default_factory=list)
    stats: Dict = field(default_factory=dict)
    errors: List[Dict] = field(default_factory=list)
    hierarchy: Dict = field(default_factory=dict)  # Tree structure of crawled pages


class DeepDocCrawler:
    """
    Deep interactive crawler for documentation sites with expandable content.
    
    Uses Playwright to:
    1. Load pages with full JS rendering
    2. Click on expandable elements (accordions, dropdowns, tree nodes)
    3. Extract revealed content and links
    4. Recursively crawl discovered pages
    """
    
    def __init__(self, config: DeepCrawlConfig = None):
        self.config = config or DeepCrawlConfig()
        self.url_normalizer = URLNormalizer()
        
        # Scope filter (initialized per-crawl in .crawl())
        self._scope_filter: Optional[ScopeFilter] = None
        
        # State
        self._visited_urls: Set[str] = set()
        self._queued_urls: Set[str] = set()   # frontier dedup — prevents double-enqueue
        self._pages: List[DeepPageData] = []
        self._errors: List[Dict] = []
        self._hierarchy: Dict = {}
        self._scope_rejected_buffer: Set[str] = set()  # buffered scope-rejected URLs for recovery
        
        # Playwright
        self._playwright = None
        self._browser: Browser = None
        self._context: BrowserContext = None
        
        # FluidTopics content API resolver
        # Populated on first page load when we detect a FluidTopics-powered
        # site (ServiceNow docs, etc.).  Maps prettyUrl → (mapId, contentId).
        self._ft_resolver: Dict[str, Tuple[str, str]] = {}  # url_path → (mapId, contentId)
        self._ft_base_url: str = ''  # e.g. "https://www.servicenow.com/docs"
        
        # Progress
        self._progress_callback: Optional[Callable] = None
        self._stop_requested = False
        
        # Stats
        self._start_time = 0
        self._expandables_clicked = 0
        self._links_discovered = 0
    
    def set_progress_callback(self, callback: Callable) -> None:
        """Set callback for progress updates: callback(pages_crawled, current_url, stats)"""
        self._progress_callback = callback
    
    def stop(self) -> None:
        """Request crawler to stop."""
        self._stop_requested = True
        logger.info("Stop requested")
    
    def _init_browser(self) -> None:
        """Initialize Playwright browser."""
        if self._playwright is None:
            self._playwright = sync_playwright().start()
            self._browser = self._playwright.chromium.launch(
                headless=self.config.headless,
                args=[
                    '--disable-gpu',
                    '--no-sandbox',
                    '--disable-dev-shm-usage',
                    '--disable-web-security',
                ]
            )
            self._context = self._browser.new_context(
                user_agent=self.config.user_agent,
                viewport={
                    'width': self.config.viewport_width,
                    'height': self.config.viewport_height
                },
                locale='en-US',
                timezone_id='America/New_York',
            )
            logger.info("Playwright browser initialized for deep crawling")
    
    def _close_browser(self) -> None:
        """Close Playwright browser."""
        if self._context:
            try:
                self._context.close()
            except Exception:
                pass
            self._context = None
        
        if self._browser:
            try:
                self._browser.close()
            except Exception:
                pass
            self._browser = None
        
        if self._playwright:
            try:
                self._playwright.stop()
            except Exception:
                pass
            self._playwright = None
    
    def _expand_all_elements(self, page: Page) -> tuple:
        """
        Delegate interactive expansion to the interaction_policy module.
        
        The policy engine handles:
        - Candidate discovery via CSS selectors + text heuristics
        - Per-click meaningful-delta gating (text / links / ARIA state)
        - Dedup fingerprinting to avoid re-clicking elements
        - Budget enforcement (max_clicks_per_page)
        
        Returns:
            Tuple of (meaningful_clicks, hit_limit)
        """
        result = interaction_policy.expansion_loop(
            page,
            max_clicks=self.config.max_clicks_per_page,
            click_timeout_ms=getattr(self.config, 'click_timeout_ms', 1500),
            delay_after_click_s=self.config.delay_after_click,
            meaningful_text_delta=getattr(self.config, 'meaningful_text_delta', 80),
            meaningful_link_delta=getattr(self.config, 'meaningful_link_delta', 1),
            selectors=self.config.interactive_selectors,
        )
        self._expandables_clicked += result.meaningful_clicks
        return result.meaningful_clicks, result.hit_limit
    
    # ------------------------------------------------------------------ #
    #  Shadow DOM & API response content extraction                        #
    # ------------------------------------------------------------------ #

    @staticmethod
    def _extract_shadow_dom_text(page: Page) -> str:
        """Recursively extract all visible text from shadow DOM trees.

        Many modern SPA documentation sites (ServiceNow / FluidTopics,
        Salesforce Lightning, etc.) render content inside Web Components
        with closed or open shadow roots.  ``page.inner_text('body')``
        only sees the light DOM — this helper reaches into every shadow
        root and collects the text content.

        Returns the concatenated text from all shadow roots, stripped of
        style/script noise.
        """
        try:
            text = page.evaluate(r"""
                () => {
                    const SKIP = new Set(['STYLE', 'SCRIPT', 'NOSCRIPT', 'SVG', 'IMG', 'BR', 'HR']);
                    function collect(root, depth) {
                        if (depth > 12) return '';
                        let out = '';
                        const nodes = root.childNodes || [];
                        for (const n of nodes) {
                            if (n.nodeType === Node.TEXT_NODE) {
                                const t = n.textContent?.trim();
                                if (t) out += t + ' ';
                            } else if (n.nodeType === Node.ELEMENT_NODE) {
                                if (SKIP.has(n.tagName)) continue;
                                if (n.shadowRoot) out += collect(n.shadowRoot, depth + 1);
                                out += collect(n, depth + 1);
                            }
                        }
                        return out;
                    }
                    // Collect from shadow DOM only (main-body text is
                    // already handled by the normal extractor)
                    let result = '';
                    document.querySelectorAll('*').forEach(el => {
                        if (el.shadowRoot) {
                            result += collect(el.shadowRoot, 0) + '\n';
                        }
                    });
                    return result.trim();
                }
            """)
            return text or ''
        except Exception as e:
            logger.debug(f"[SHADOW-DOM] Extraction failed: {e}")
            return ''

    @staticmethod
    def _parse_intercepted_html(html: str) -> str:
        """Parse intercepted HTML body into clean text."""
        try:
            soup = BeautifulSoup(html, _BS_PARSER)
            # Remove scripts/styles
            for tag in soup(['script', 'style', 'noscript']):
                tag.decompose()
            text = soup.get_text(separator=' ', strip=True)
            # Collapse whitespace
            text = re.sub(r'\n{3,}', '\n\n', text)
            text = re.sub(r' {2,}', ' ', text)
            return text
        except Exception:
            return ''

    # ------------------------------------------------------------------ #
    #  FluidTopics content resolver                                        #
    # ------------------------------------------------------------------ #

    def _ft_build_resolver(self, pages_json: dict, map_id: str, base_url: str) -> None:
        """Build a prettyUrl → (mapId, contentId) mapping from FluidTopics pages API.

        Called when we intercept the ``/api/khub/maps/{mapId}/pages`` response.
        The resulting resolver lets us fetch topic content directly via HTTP
        for pages where the SPA doesn't render.
        """
        self._ft_base_url = base_url.rstrip('/')

        def _walk(node: dict) -> None:
            pretty = node.get('prettyUrl', '')
            content_id = node.get('contentId', '')
            if pretty and content_id:
                self._ft_resolver[pretty] = (map_id, content_id)
            for child in node.get('children', []):
                _walk(child)
            for child in node.get('pageToc', []):
                _walk(child)

        # paginatedToc is the top-level list of root pages
        for root in pages_json.get('paginatedToc', []):
            _walk(root)

        logger.info(
            f"[FT-RESOLVER] Built URL→content map: "
            f"{len(self._ft_resolver)} topics from map {map_id}"
        )

    def _ft_fetch_content(self, url: str) -> str:
        """Fetch topic content from FluidTopics API for the given page URL.

        Returns raw HTML or empty string if not resolvable.
        """
        if not self._ft_resolver:
            return ''

        # Extract the /r/... path portion from the URL
        parsed = urlparse(url)
        path = parsed.path  # e.g. /docs/r/it-asset-management/.../page.html

        # Try the path directly, then strip /docs prefix
        for candidate in [path, re.sub(r'^/docs', '', path)]:
            if candidate in self._ft_resolver:
                map_id, content_id = self._ft_resolver[candidate]
                api_url = (
                    f"{self._ft_base_url}/api/khub/maps/{map_id}"
                    f"/topics/{content_id}/content"
                )
                try:
                    resp = requests.get(api_url, timeout=10)
                    if resp.status_code == 200 and len(resp.text) > 100:
                        logger.info(
                            f"[FT-FETCH] Got {len(resp.text):,} bytes "
                            f"for {candidate[:60]}"
                        )
                        return resp.text
                except Exception as e:
                    logger.debug(f"[FT-FETCH] Error for {candidate}: {e}")
                break

        return ''

    def _dismiss_cookie_consent(self, page: Page) -> None:
        """Try to dismiss cookie consent banners that block content."""
        consent_selectors = [
            # Common consent button patterns
            'button:has-text("Accept")',
            'button:has-text("Accept and Proceed")',
            'button:has-text("Accept All")',
            'button:has-text("Accept Cookies")',
            'button:has-text("I Accept")',
            'button:has-text("Agree")',
            'button:has-text("OK")',
            'button:has-text("Got it")',
            'button:has-text("Allow All")',
            # Common class/id patterns
            '#onetrust-accept-btn-handler',
            '.cookie-accept',
            '.cc-accept',
            '[data-testid="cookie-accept"]',
            '.consent-accept',
            '#cookie-accept',
            '#accept-cookies',
            '.js-accept-cookies',
        ]
        for selector in consent_selectors:
            try:
                btn = page.query_selector(selector)
                if btn and btn.is_visible():
                    btn.click(timeout=2000)
                    logger.debug(f"[COOKIE] Dismissed consent banner via: {selector}")
                    page.wait_for_timeout(500)
                    return
            except Exception:
                continue

    def _wait_for_spa_content(self, page: Page) -> None:
        """Wait for SPA content to load if we detect a loading placeholder.

        Checks if the visible text is just a loading message or cookie banner,
        and if so, waits up to 10 additional seconds for real content to appear.
        """
        try:
            body_text = page.evaluate("""
                () => (document.body?.innerText || '').trim().substring(0, 500)
            """)
            # Detect loading/placeholder patterns
            loading_patterns = [
                'loading application',
                'loading...',
                'please wait',
                'initializing',
                'just a moment',
            ]
            is_loading = any(pat in body_text.lower() for pat in loading_patterns)
            is_mostly_cookie = (
                'cookie' in body_text.lower()
                and len(body_text) < 600
                and 'accept' in body_text.lower()
            )

            if is_loading or is_mostly_cookie:
                logger.info(f"[SPA-WAIT] Detected loading/placeholder — waiting for real content...")
                # Wait for content selectors to appear
                content_appeared = False
                for selector in ['main', 'article', '.content', '#content', '[role="main"]',
                                 '.documentation', '.doc-content', '.markdown-section']:
                    try:
                        page.wait_for_selector(selector, timeout=8000, state='attached')
                        content_appeared = True
                        logger.debug(f"[SPA-WAIT] Content appeared via: {selector}")
                        break
                    except PlaywrightTimeout:
                        continue

                if not content_appeared:
                    # Fallback: just wait and hope
                    page.wait_for_timeout(5000)
                    logger.debug("[SPA-WAIT] No content selector appeared, waited 5s")
                else:
                    # Give a bit more time for content to render
                    page.wait_for_timeout(1500)
        except Exception as e:
            logger.debug(f"[SPA-WAIT] Error: {e}")

    def _detect_page_complexity(self, page: Page) -> str:
        """Auto-detect whether a page needs JS expansion or is simple HTML.

        Runs quickly inside Playwright after page load.  Checks two things:
          1. Are there interactive expandable elements on the page?
          2. Did JS rendering actually produce meaningful content?

        Returns:
            'html'   — static page, skip expansion (fast path)
            'js'     — JS-heavy page, do full expansion
        """
        try:
            # Count expandable / interactive elements on the page
            expandable_count = page.evaluate("""
                () => {
                    const selectors = [
                        '[aria-expanded="false"]',
                        'details:not([open]) > summary',
                        '.collapsed',
                        '.accordion-button.collapsed',
                        '[data-toggle]',
                        '[data-bs-toggle]',
                        '[class*="expand"]:not([class*="expanded"])',
                        '[class*="collapse"]:not([class*="collapsed"])',
                        '[role="tab"]:not([aria-selected="true"])',
                        '.load-more', '.show-more', '.view-more',
                        '.tree-node:not(.expanded)',
                        '.toc-item > .toc-link',
                        '.ohc-sidebar-item',
                    ];
                    let count = 0;
                    for (const sel of selectors) {
                        try { count += document.querySelectorAll(sel).length; }
                        catch(e) {}
                    }
                    return count;
                }
            """)

            # Count visible <a> links already on the page
            link_count = page.evaluate("""
                () => document.querySelectorAll('a[href]').length
            """)

            # Heuristic decision:
            # - If there are expandable elements -> JS mode (needs expansion)
            # - If very few links found -> JS mode (content may be loading)
            # - Otherwise -> HTML mode (fast path)
            if expandable_count >= 3:
                logger.info(
                    f"[AUTO-DETECT] JS mode — {expandable_count} expandable "
                    f"elements found (links={link_count})"
                )
                return 'js'

            if link_count < 5:
                # Very few links might mean JS hasn't loaded navigation yet
                # but also might be a legitimate leaf page -- check body text
                body_text_len = page.evaluate("""
                    () => (document.body?.innerText || '').trim().length
                """)
                if body_text_len < 200:
                    logger.info(
                        f"[AUTO-DETECT] JS mode — sparse page "
                        f"(links={link_count}, text={body_text_len} chars)"
                    )
                    return 'js'

            logger.info(
                f"[AUTO-DETECT] HTML mode — static page "
                f"(links={link_count}, expandables={expandable_count})"
            )
            return 'html'

        except Exception as e:
            logger.debug(f"[AUTO-DETECT] Error during detection: {e} — defaulting to js")
            return 'js'

    def _extract_links_from_page(self, page: Page, base_url: str) -> List[str]:
        """Extract all internal, in-scope links from the page after expansion."""
        links = set()
        base_domain = urlparse(base_url).netloc.lower()
        total_hrefs = 0
        rejected_external = 0
        rejected_normalize = 0
        rejected_scope = 0
        rejected_nonhttp = 0
        rejected_dedup = 0
        rejected_junk = 0

        # Compile built-in deny patterns for junk URL filtering
        junk_patterns = []
        for pat in self.config.builtin_deny_patterns:
            try:
                junk_patterns.append(re.compile(pat, re.IGNORECASE))
            except re.error:
                pass
        
        for selector in self.config.link_selectors:
            try:
                elements = page.query_selector_all(selector)
                
                for element in elements:
                    try:
                        href = element.get_attribute('href')
                        if not href:
                            continue
                        total_hrefs += 1
                        
                        # Resolve relative URLs (ensure directory paths keep trailing /)
                        absolute_url = urljoin(ensure_joinable_base(page.url), href)
                        
                        # Skip non-http links
                        if not absolute_url.startswith(('http://', 'https://')):
                            rejected_nonhttp += 1
                            continue
                        
                        # Skip external links (fast domain check before full scope filter)
                        url_domain = urlparse(absolute_url).netloc.lower()
                        if url_domain != base_domain:
                            rejected_external += 1
                            continue
                        
                        # Strip fragments
                        if '#' in absolute_url:
                            absolute_url = absolute_url.split('#')[0]
                        
                        # Normalize
                        normalized = self.url_normalizer.normalize(absolute_url)
                        if not normalized:
                            rejected_normalize += 1
                            continue

                        # Junk URL filter — viewer, attachment, locale variants
                        if any(rx.search(normalized) for rx in junk_patterns):
                            rejected_junk += 1
                            continue

                        # Scope gate — reject outside-subtree links here,
                        # before they ever reach the BFS queue
                        if self._scope_filter and not self._scope_filter.accept(normalized):
                            rejected_scope += 1
                            self._scope_rejected_buffer.add(normalized)
                            continue

                        if normalized in links:
                            rejected_dedup += 1
                        links.add(normalized)
                            
                    except Exception:
                        continue
                        
            except Exception:
                continue
        
        self._links_discovered += len(links)
        logger.info(
            f"[LINKS] {page.url[:60]} → "
            f"hrefs_found={total_hrefs} "
            f"in_scope={len(links)} "
            f"external={rejected_external} "
            f"nonhttp={rejected_nonhttp} "
            f"bad_normalize={rejected_normalize} "
            f"scope_rejected={rejected_scope} "
            f"junk_filtered={rejected_junk} "
            f"dedup={rejected_dedup}"
        )
        return list(links)
    
    def _extract_breadcrumb(self, page: Page) -> List[str]:
        """Extract breadcrumb navigation."""
        breadcrumb = []
        
        selectors = [
            '.breadcrumb a',
            '.breadcrumb li',
            '[aria-label="breadcrumb"] a',
            '.ohc-breadcrumb a',
            'nav[aria-label*="breadcrumb"] a',
        ]
        
        for selector in selectors:
            try:
                elements = page.query_selector_all(selector)
                if elements:
                    for el in elements:
                        text = el.inner_text().strip()
                        if text and text not in breadcrumb:
                            breadcrumb.append(text)
                    if breadcrumb:
                        break
            except Exception:
                continue
        
        return breadcrumb
    
    def _extract_section_path(self, page: Page) -> List[str]:
        """Extract the hierarchical section path from sidebar/TOC."""
        path = []
        
        selectors = [
            '.toc-item.active',
            '.nav-item.active',
            '.tree-item.selected',
            '[aria-current="page"]',
            '.ohc-sidebar-item.active',
            '.is-selected',
        ]
        
        for selector in selectors:
            try:
                # Find the active item and its ancestors
                active = page.query_selector(selector)
                if active:
                    # Try to get parent items
                    current = active
                    while current:
                        text = current.inner_text().strip().split('\n')[0][:100]
                        if text:
                            path.insert(0, text)
                        
                        # Move to parent
                        parent = current.evaluate("""el => {
                            const parent = el.closest('li, .toc-item, .tree-item, .nav-item');
                            if (parent && parent.parentElement) {
                                const grandparent = parent.parentElement.closest('li, .toc-item, .tree-item, .nav-item');
                                return grandparent ? true : false;
                            }
                            return false;
                        }""")
                        
                        if not parent:
                            break
                        
                        current = current.evaluate_handle("""el => {
                            const parent = el.closest('li, .toc-item, .tree-item, .nav-item');
                            if (parent && parent.parentElement) {
                                return parent.parentElement.closest('li, .toc-item, .tree-item, .nav-item');
                            }
                            return null;
                        }""").as_element()
                        
                        if not current:
                            break
                    
                    if path:
                        break
            except Exception:
                continue
        
        return path
    
    def _extract_content(
        self,
        page: Page,
        intercepted_html: str = '',
        page_url: str = '',
    ) -> Dict[str, Any]:
        """Extract main content from the page.

        Extraction strategy (in order of preference):
        1. Normal DOM extraction (light DOM content selectors)
        2. Shadow DOM recursive text extraction (Web Components)
        3. Intercepted API response HTML (FluidTopics, etc.)
        """
        content = {
            'title': '',
            'headings': {},
            'text': '',
            'tables': [],
            'code_blocks': [],
        }
        
        # Extract title
        try:
            title_el = page.query_selector('h1') or page.query_selector('title')
            if title_el:
                content['title'] = title_el.inner_text().strip()
        except Exception:
            pass
        
        # Find main content area
        main_content = None
        for selector in self.config.content_selectors:
            try:
                main_content = page.query_selector(selector)
                if main_content:
                    break
            except Exception:
                continue
        
        if not main_content:
            main_content = page.query_selector('body')
        
        if main_content:
            # Extract headings
            for level in range(1, 7):
                try:
                    headings = main_content.query_selector_all(f'h{level}')
                    if headings:
                        content['headings'][f'h{level}'] = [
                            h.inner_text().strip() for h in headings if h.inner_text().strip()
                        ]
                except Exception:
                    pass
            
            # Extract tables
            try:
                tables = main_content.query_selector_all('table')
                for table in tables:
                    try:
                        # Get table as structured data
                        table_data = {
                            'headers': [],
                            'rows': []
                        }
                        
                        # Headers
                        headers = table.query_selector_all('th')
                        table_data['headers'] = [h.inner_text().strip() for h in headers]
                        
                        # Rows
                        rows = table.query_selector_all('tr')
                        for row in rows:
                            cells = row.query_selector_all('td')
                            if cells:
                                table_data['rows'].append([c.inner_text().strip() for c in cells])
                        
                        if table_data['headers'] or table_data['rows']:
                            content['tables'].append(table_data)
                    except Exception:
                        pass
            except Exception:
                pass
            
            # Extract code blocks
            try:
                code_blocks = main_content.query_selector_all('pre, code')
                for block in code_blocks:
                    try:
                        code_text = block.inner_text().strip()
                        if code_text and len(code_text) > 10:
                            content['code_blocks'].append(code_text)
                    except Exception:
                        pass
            except Exception:
                pass
            
            # Extract text content
            try:
                # Remove excluded elements
                for selector in self.config.exclude_selectors:
                    try:
                        for el in main_content.query_selector_all(selector):
                            el.evaluate("el => el.remove()")
                    except Exception:
                        pass
                
                content['text'] = main_content.inner_text().strip()
                # Clean up whitespace
                content['text'] = re.sub(r'\n{3,}', '\n\n', content['text'])
                content['text'] = re.sub(r' {2,}', ' ', content['text'])
            except Exception:
                pass
        
        # ── Fallback 1: Shadow DOM text extraction ──────────────────
        # If the normal DOM extraction returned little/no content,
        # try recursively extracting from shadow roots (Web Components).
        dom_word_count = len(content['text'].split())
        if dom_word_count < self.config.min_word_count:
            shadow_text = self._extract_shadow_dom_text(page)
            shadow_words = len(shadow_text.split())
            if shadow_words > dom_word_count and shadow_words >= self.config.min_word_count:
                logger.info(
                    f"[SHADOW-DOM] Using shadow DOM text: {shadow_words:,} words "
                    f"(light DOM had {dom_word_count})"
                )
                content['text'] = shadow_text

        # ── Fallback 2: Intercepted API response ────────────────────
        # Many SPA doc sites (FluidTopics / ServiceNow, etc.) fetch the
        # real page content via an internal API and render it inside
        # shadow DOM or Web Components that may not fully hydrate in
        # headless mode.  If we captured a content-rich HTML response
        # during page load, prefer it when it has substantially more
        # text than what DOM + shadow-DOM extraction produced.
        current_word_count = len(content['text'].split())
        if intercepted_html:
            api_text = self._parse_intercepted_html(intercepted_html)
            api_words = len(api_text.split())
            # Use API content if:
            #   a) current content is below threshold AND API has enough, OR
            #   b) API content is ≥3× richer (i.e. DOM/shadow got nav junk)
            use_api = (
                (current_word_count < self.config.min_word_count
                 and api_words >= self.config.min_word_count)
                or (api_words >= 3 * max(current_word_count, 1))
            )
            if use_api:
                logger.info(
                    f"[API-CONTENT] Using intercepted API response: "
                    f"{api_words:,} words (DOM/shadow had {current_word_count})"
                )
                content['text'] = api_text
                # Also extract structured data from the API HTML
                try:
                    soup = BeautifulSoup(intercepted_html, _BS_PARSER)
                    # Title
                    if not content['title'] or content['title'] == page.url:
                        h1 = soup.find('h1')
                        if h1:
                            content['title'] = h1.get_text(strip=True)
                    # Headings
                    for level in range(1, 7):
                        hs = soup.find_all(f'h{level}')
                        if hs:
                            content['headings'][f'h{level}'] = [
                                h.get_text(strip=True) for h in hs
                                if h.get_text(strip=True)
                            ]
                    # Tables
                    if not content['tables']:
                        for table in soup.find_all('table'):
                            table_data = {'headers': [], 'rows': []}
                            for th in table.find_all('th'):
                                table_data['headers'].append(th.get_text(strip=True))
                            for tr in table.find_all('tr'):
                                cells = tr.find_all('td')
                                if cells:
                                    table_data['rows'].append(
                                        [c.get_text(strip=True) for c in cells]
                                    )
                            if table_data['headers'] or table_data['rows']:
                                content['tables'].append(table_data)
                    # Code blocks
                    if not content['code_blocks']:
                        for block in soup.find_all(['pre', 'code']):
                            code_text = block.get_text(strip=True)
                            if code_text and len(code_text) > 10:
                                content['code_blocks'].append(code_text)
                except Exception:
                    pass
        
        # ── Fallback 3: FluidTopics direct content fetch ────────────
        # If we detected a FluidTopics-powered site (resolver built from
        # /api/khub/maps/.../pages), fetch topic content directly via
        # HTTP using the URL→contentId mapping.  Always try when the
        # resolver is available and current content is thin (<300 words),
        # since shadow-DOM extraction often captures only nav/header text.
        current_word_count = len(content['text'].split())
        if self._ft_resolver and current_word_count < 300:
            ft_html = self._ft_fetch_content(page_url or page.url)
            if ft_html:
                ft_text = self._parse_intercepted_html(ft_html)
                ft_words = len(ft_text.split())
                use_ft = (
                    ft_words >= self.config.min_word_count
                    and ft_words > current_word_count
                )
                if use_ft:
                    logger.info(
                        f"[FT-FETCH] Using FluidTopics API content: "
                        f"{ft_words:,} words (previous had {current_word_count})"
                    )
                    content['text'] = ft_text
                    # Extract structured data from FT HTML
                    try:
                        soup = BeautifulSoup(ft_html, _BS_PARSER)
                        if not content['title']:
                            h1 = soup.find('h1')
                            if h1:
                                content['title'] = h1.get_text(strip=True)
                        for level in range(1, 7):
                            hs = soup.find_all(f'h{level}')
                            if hs:
                                content['headings'][f'h{level}'] = [
                                    h.get_text(strip=True) for h in hs
                                    if h.get_text(strip=True)
                                ]
                        if not content['tables']:
                            for table in soup.find_all('table'):
                                td = {'headers': [], 'rows': []}
                                for th in table.find_all('th'):
                                    td['headers'].append(th.get_text(strip=True))
                                for tr in table.find_all('tr'):
                                    cells = tr.find_all('td')
                                    if cells:
                                        td['rows'].append(
                                            [c.get_text(strip=True) for c in cells]
                                        )
                                if td['headers'] or td['rows']:
                                    content['tables'].append(td)
                        if not content['code_blocks']:
                            for block in soup.find_all(['pre', 'code']):
                                code_text = block.get_text(strip=True)
                                if code_text and len(code_text) > 10:
                                    content['code_blocks'].append(code_text)
                    except Exception:
                        pass

        # Per-page scrape summary
        heading_count = sum(len(v) for v in content['headings'].values())
        text_len = len(content['text'])
        logger.info(
            f"[SCRAPE] title='{content['title'][:60]}' | "
            f"text={text_len:,} chars | "
            f"headings={heading_count} | "
            f"tables={len(content['tables'])} | "
            f"code_blocks={len(content['code_blocks'])}"
        )
        
        return content
    
    def _crawl_page(
        self,
        url: str,
        depth: int,
        parent_url: str,
        section_path: List[str],
        base_url: str
    ) -> Optional[DeepPageData]:
        """Crawl a single page with expansion."""
        
        # Check limits
        if len(self._pages) >= self.config.max_pages:
            return None
        
        if depth > self.config.max_depth:
            return None
        
        if self._stop_requested:
            return None
        
        # Normalize URL
        normalized_url = self.url_normalizer.normalize(url)
        if not normalized_url:
            return None
        
        # Check if already visited
        if normalized_url in self._visited_urls:
            return None
        
        # Scope guard (belt-and-suspenders — primary filtering is at link extraction)
        if self._scope_filter and not self._scope_filter.accept(normalized_url):
            logger.debug(f"[SCOPE] Rejected at crawl-page gate: {normalized_url}")
            return None
        
        # Mark visited AFTER scope check so scope-rejected URLs aren't permanently eaten
        self._visited_urls.add(normalized_url)
        
        logger.info(f"Deep crawling [{depth}]: {normalized_url}")
        
        try:
            # Create new page
            page = self._context.new_page()
            
            # ── Response interception for SPA content APIs ──────────
            # Web-component SPAs (FluidTopics, etc.) often fetch page
            # content via internal APIs (JSON-wrapped HTML, or raw HTML)
            # then render it inside shadow DOM — invisible to normal
            # DOM extraction.  We capture the largest HTML-bearing
            # response so we can fall back to it.
            #
            # We also detect FluidTopics sites and build a URL→content
            # resolver from the /api/khub/maps/{id}/pages response.
            _intercepted = {'html': '', 'size': 0}
            _self = self  # capture for closure

            def _on_response(resp):
                try:
                    ct = resp.headers.get('content-type', '')
                    rurl = resp.url
                    url_lower = rurl.lower()

                    # 1. Capture largest HTML response from content APIs
                    if (
                        'text/html' in ct
                        and '/api/' in url_lower
                        and resp.status == 200
                    ):
                        body = resp.text()
                        if len(body) > _intercepted['size']:
                            _intercepted['html'] = body
                            _intercepted['size'] = len(body)

                    # 2. Detect FluidTopics pages API and build resolver
                    #    Pattern: /api/khub/maps/{mapId}/pages
                    if (
                        'application/json' in ct
                        and resp.status == 200
                        and '/api/khub/maps/' in rurl
                        and '/pages' in rurl
                        and not _self._ft_resolver
                    ):
                        import re as _re
                        m = _re.search(r'/api/khub/maps/([^/]+)/pages', rurl)
                        if m:
                            map_id = m.group(1)
                            parsed_origin = urlparse(rurl)
                            ft_base = f"{parsed_origin.scheme}://{parsed_origin.netloc}/docs"
                            try:
                                body = resp.json()
                                _self._ft_build_resolver(body, map_id, ft_base)
                            except Exception:
                                pass
                except Exception:
                    pass

            page.on('response', _on_response)
            
            try:
                # Navigate to URL
                response = page.goto(
                    normalized_url,
                    timeout=self.config.timeout,
                    wait_until='load'
                )
                
                if response is None or response.status >= 400:
                    self._errors.append({
                        'url': normalized_url,
                        'error': f"HTTP {response.status if response else 'No response'}",
                        'depth': depth
                    })
                    return None
                
                # Wait for content - USE domcontentloaded, NOT networkidle
                page.wait_for_load_state('domcontentloaded')
                
                # Wait for dynamic content (JS-rendered TOC / sidebars).
                # Many doc sites (Oracle, etc.) load navigation via RequireJS
                # after DOMContentLoaded.  We poll for <a href> stabilisation
                # within the existing per-page timeout budget.
                try:
                    page.wait_for_load_state('networkidle', timeout=5000)
                except PlaywrightTimeout:
                    pass   # best-effort — don't fail the page on this
                
                # Brief stabilization wait
                page.wait_for_timeout(1000)
                
                # ── Dismiss cookie consent banners ───────────────
                # Many sites (ServiceNow, etc.) block content behind
                # cookie consent overlays.  Click common accept buttons.
                self._dismiss_cookie_consent(page)
                
                # ── Wait for SPA content to load ────────────────
                # Sites like ServiceNow render via heavy JS frameworks.
                # If we detect "Loading" placeholder text, wait longer
                # for the real content to appear.
                self._wait_for_spa_content(page)
                
                # ── Redirect detection: widen scope if needed ───────
                # If Playwright followed a redirect (JS/server), the landing
                # URL may be outside the original scope.  Widen the scope to
                # the common path prefix of the original and landing URLs so
                # links on the landing page are not all scope-rejected.
                landing_url = page.url
                if landing_url and landing_url != normalized_url:
                    if self._scope_filter:
                        widened = self._scope_filter.widen_scope(landing_url)
                        if widened:
                            logger.info(
                                f"[SCOPE] Redirect detected: "
                                f"{normalized_url[:50]} → {landing_url[:50]}"
                            )
                            self._scope_filter.log_scope()
                
                # ── Auto-detect page complexity ─────────────────────
                # Determine if this page needs JS expansion or is simple HTML.
                # This replaces the old manual standard/deep mode choice.
                page_type = self._detect_page_complexity(page)
                
                # Skip interactive expansion when the BFS queue already has
                # plenty of URLs.  The sidebar / TOC is typically identical on
                # every page, so re-expanding it on every visit is pure waste.
                queue_size = len(self._queue)
                if queue_size >= self.config.max_pages:
                    logger.debug(
                        f"Skipping expansion — queue already has {queue_size} URLs"
                    )
                    expanded, hit_limit = 0, False
                elif page_type == 'html':
                    # HTML page — skip expansion entirely (fast path)
                    logger.info(f"[FAST] Skipping expansion for HTML page: {normalized_url[:60]}")
                    expanded, hit_limit = 0, False
                else:
                    # JS-heavy page — do full expansion
                    logger.info(f"[DEEP] Running expansion for JS page: {normalized_url[:60]}")
                    expanded, hit_limit = self._expand_all_elements(page)
                    if hit_limit:
                        logger.info(f"Expansion limit reached on {normalized_url}")
                    else:
                        logger.debug(f"Expanded {expanded} elements on {normalized_url}")
                
                # Brief wait after expansion (only if we expanded something)
                if expanded > 0:
                    page.wait_for_timeout(500)
                
                # Extract content (pass intercepted API HTML for fallback)
                content = self._extract_content(
                    page,
                    intercepted_html=_intercepted['html'],
                    page_url=normalized_url,
                )
                
                # Extract links — skip heavy extraction when queue is already
                # larger than max_pages (sidebar links are the same on every
                # page; re-extracting 22K hrefs per page is pure waste).
                if len(self._queue) >= self.config.max_pages:
                    links = []
                    logger.debug(
                        f"Skipping link extraction — queue has {len(self._queue)} URLs"
                    )
                else:
                    links = self._extract_links_from_page(page, base_url)
                
                # Extract breadcrumb and section path
                breadcrumb = self._extract_breadcrumb(page)
                current_section = self._extract_section_path(page) or section_path
                
                # Build page data
                page_data = DeepPageData(
                    url=normalized_url,
                    title=content['title'],
                    breadcrumb=breadcrumb,
                    section_path=current_section,
                    headings=content['headings'],
                    text_content=content['text'],
                    tables=content['tables'],
                    code_blocks=content['code_blocks'],
                    internal_links=links,
                    parent_url=parent_url,
                    depth=depth,
                    word_count=len(content['text'].split())
                )
                
                logger.info(
                    f"[PAGE OK] {normalized_url[:70]} — "
                    f"{page_data.word_count:,} words, "
                    f"{len(links)} links, "
                    f"depth={depth}"
                )
                
                # ── Content quality gate ────────────────────────────
                # Skip pages that are empty, just cookie banners, or
                # "Loading application..." SPA placeholders.
                if page_data.word_count < self.config.min_word_count:
                    # Check if it's a loading / cookie-only page
                    text_lower = page_data.text_content.lower()
                    is_junk = (
                        'loading application' in text_lower
                        or ('cookie' in text_lower and page_data.word_count < 100)
                        or page_data.word_count == 0
                    )
                    if is_junk:
                        logger.warning(
                            f"[SKIP] {normalized_url[:70]} — "
                            f"empty/loading/cookie page ({page_data.word_count} words)"
                        )
                        # Still return links so the frontier keeps growing,
                        # but mark page as skipped so it's not counted
                        # toward the page limit
                        page_data.text_content = ""
                        page_data.word_count = 0
                        page_data._skipped = True
                        return page_data
                
                # Progress callback
                if self._progress_callback:
                    try:
                        stats = {
                            'pages_crawled': len(self._pages) + 1,
                            'expandables_clicked': self._expandables_clicked,
                            'links_discovered': self._links_discovered,
                        }
                        self._progress_callback(len(self._pages) + 1, normalized_url, stats)
                    except Exception:
                        pass
                
                return page_data
                
            finally:
                page.close()
                
        except PlaywrightTimeout:
            logger.warning(
                f"[PAGE-FALLBACK] Timeout on {normalized_url} — "
                f"falling back to static extraction (queue continues)"
            )
            if self.config.enable_static_fallback:
                return self._crawl_page_static(normalized_url, depth, parent_url, section_path, base_url)
            self._errors.append({
                'url': normalized_url,
                'error': 'Timeout (no static fallback)',
                'depth': depth,
                'fallback': False,
            })
            return None
        except Exception as e:
            logger.warning(
                f"[PAGE-FALLBACK] Deep-mode error on {normalized_url}: {e} — "
                f"falling back to static extraction (queue continues)"
            )
            if self.config.enable_static_fallback:
                return self._crawl_page_static(normalized_url, depth, parent_url, section_path, base_url)
            self._errors.append({
                'url': normalized_url,
                'error': str(e),
                'depth': depth,
                'fallback': False,
            })
            return None
    
    def _crawl_page_static(
        self,
        url: str,
        depth: int,
        parent_url: str,
        section_path: List[str],
        base_url: str
    ) -> Optional[DeepPageData]:
        """
        Static fallback crawler using requests + BeautifulSoup.
        
        Invoked **per-page** when Playwright fails (timeout, JS errors, etc.)
        The crawl queue is NOT affected — only this single URL uses static mode.
        """
        logger.info(f"[STATIC-FALLBACK] Extracting [{depth}]: {url}")
        
        try:
            headers = {
                'User-Agent': self.config.user_agent,
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
                'Accept-Language': 'en-US,en;q=0.9',
            }
            
            response = requests.get(url, headers=headers, timeout=15)
            
            if response.status_code >= 400:
                self._errors.append({
                    'url': url,
                    'error': f"HTTP {response.status_code} (static fallback)",
                    'depth': depth
                })
                return None
            
            html = response.text
            soup = BeautifulSoup(html, _BS_PARSER)
            
            # Extract title
            title = ""
            title_tag = soup.find('title')
            if title_tag:
                title = title_tag.get_text(strip=True)
            elif soup.find('h1'):
                title = soup.find('h1').get_text(strip=True)
            
            # Extract headings
            headings = {}
            for level in range(1, 7):
                h_tags = soup.find_all(f'h{level}')
                if h_tags:
                    headings[f'h{level}'] = [h.get_text(strip=True) for h in h_tags if h.get_text(strip=True)]
            
            # Extract text content from main content area
            main_content = None
            for selector in ['main', 'article', '.content', '#content', '[role="main"]']:
                main_content = soup.select_one(selector)
                if main_content:
                    break
            
            if not main_content:
                main_content = soup.find('body')
            
            text_content = ""
            if main_content:
                # Remove script, style, nav, etc.
                for tag in main_content.find_all(['script', 'style', 'nav', 'footer', 'header', 'aside']):
                    tag.decompose()
                text_content = main_content.get_text(separator=' ', strip=True)
                text_content = re.sub(r'\s+', ' ', text_content)
            
            # Extract internal links (scope-filtered)
            base_domain = urlparse(base_url).netloc.lower()
            internal_links = []
            total_hrefs = 0
            rejected_scope = 0
            for a_tag in soup.find_all('a', href=True):
                href = a_tag['href']
                abs_url = urljoin(ensure_joinable_base(url), href)
                if abs_url.startswith(('http://', 'https://')):
                    link_domain = urlparse(abs_url).netloc.lower()
                    if link_domain == base_domain:
                        total_hrefs += 1
                        # Remove fragment
                        if '#' in abs_url:
                            abs_url = abs_url.split('#')[0]
                        normalized = self.url_normalizer.normalize(abs_url)
                        if not normalized or normalized in internal_links:
                            continue
                        # Scope gate
                        if self._scope_filter and not self._scope_filter.accept(normalized):
                            rejected_scope += 1
                            continue
                        internal_links.append(normalized)
            
            self._links_discovered += len(internal_links)
            logger.info(
                f"[LINKS-STATIC] {url[:60]} → "
                f"hrefs_found={total_hrefs} "
                f"in_scope={len(internal_links)} "
                f"scope_rejected={rejected_scope}"
            )
            
            page_data = DeepPageData(
                url=url,
                title=title,
                breadcrumb=[],
                section_path=section_path,
                headings=headings,
                text_content=text_content[:50000],  # Limit content size
                tables=[],
                code_blocks=[],
                internal_links=internal_links,
                parent_url=parent_url,
                depth=depth,
                word_count=len(text_content.split())
            )
            
            return page_data
            
        except Exception as e:
            logger.warning(f"[STATIC-FALLBACK] Also failed for {url}: {e}")
            self._errors.append({
                'url': url,
                'error': f"Static fallback also failed: {str(e)}",
                'depth': depth,
                'fallback': True,
            })
            return None
    
    def crawl(self, start_url: str) -> DeepCrawlResult:
        """
        Start deep crawling from the given URL.
        
        Args:
            start_url: Starting URL (documentation index page)
            
        Returns:
            DeepCrawlResult with all crawled pages
        """
        # Initialize scope filter for this crawl
        self._scope_filter = ScopeFilter(
            root_url=start_url,
            deny_patterns=self.config.deny_patterns,
            strip_all_queries=self.config.strip_all_queries,
        )
        self._scope_filter.log_scope()
        scope_desc = self._scope_filter.scope_description
        
        logger.info("=" * 60)
        logger.info(f"DEEP CRAWL STARTED")
        logger.info(f"Start URL: {start_url}")
        logger.info(f"Scope: {scope_desc}")
        logger.info(f"Limits: max_pages={self.config.max_pages}, max_depth={self.config.max_depth}")
        logger.info(f"Per-page: timeout={self.config.timeout}ms, max_clicks={self.config.max_clicks_per_page}")
        logger.info("=" * 60)
        
        # Reset state
        self._visited_urls.clear()
        self._queued_urls.clear()
        self._pages.clear()
        self._errors.clear()
        self._hierarchy.clear()
        self._scope_rejected_buffer.clear()
        self._stop_requested = False
        self._expandables_clicked = 0
        self._links_discovered = 0
        self._start_time = time.time()
        stop_reason = "completed"
        
        # Initialize browser
        self._init_browser()
        
        try:
            # Queue: (url, depth, parent_url, section_path)
            queue = [(start_url, 0, "", [])]
            self._queue = queue   # expose to _crawl_page for skip-expansion logic
            self._queued_urls.add(start_url)
            
            while queue and not self._stop_requested:
                # Check page limit (only count non-skipped pages)
                good_count = sum(1 for p in self._pages if not getattr(p, '_skipped', False))
                if good_count >= self.config.max_pages:
                    stop_reason = f"MAX_PAGES limit reached ({self.config.max_pages})"
                    logger.info(f"STOPPING: {stop_reason}")
                    break
                
                url, depth, parent_url, section_path = queue.pop(0)
                
                # Check depth limit
                if depth > self.config.max_depth:
                    logger.debug(f"Skipping {url} - exceeds max depth {self.config.max_depth}")
                    continue
                
                # Log progress
                elapsed = time.time() - self._start_time
                logger.info(f"[{good_count+1}/{self.config.max_pages}] Depth:{depth} | {elapsed:.1f}s | Queue:{len(queue)} | {url[:80]}...")
                
                # Crawl page
                page_data = self._crawl_page(url, depth, parent_url, section_path, start_url)
                
                if page_data:
                    self._pages.append(page_data)
                    
                    # Add discovered links to global frontier
                    if depth < self.config.max_depth:
                        new_enqueued = 0
                        rejected_scope = 0
                        rejected_visited = 0
                        for link in page_data.internal_links:
                            # Skip already-visited or already-queued URLs
                            if link in self._visited_urls:
                                rejected_visited += 1
                                continue
                            if link in self._queued_urls:
                                continue
                            # Links are already scope-filtered at extraction time,
                            # but belt-and-suspenders check here too
                            if self._scope_filter and not self._scope_filter.accept(link):
                                rejected_scope += 1
                                continue
                            queue.append((
                                link,
                                depth + 1,
                                page_data.url,
                                page_data.section_path.copy()
                            ))
                            self._queued_urls.add(link)
                            new_enqueued += 1
                        
                        # Use info for pages that actually discover links,
                        # debug for pages where extraction was skipped
                        log_fn = logger.info if len(page_data.internal_links) > 0 else logger.debug
                        log_fn(
                            f"[FRONTIER] Page {page_data.url[:60]} → "
                            f"discovered={len(page_data.internal_links)} "
                            f"enqueued={new_enqueued} "
                            f"already_visited={rejected_visited} "
                            f"scope_rejected={rejected_scope} "
                            f"queue_size={len(queue)}"
                        )
                        
                        # ── Zero-links fallback: widen scope if first page
                        # found no in-scope links but had scope-rejected ones ──
                        if (len(self._pages) == 1
                                and new_enqueued == 0
                                and len(self._scope_rejected_buffer) > 0
                                and self._scope_filter
                                and self._scope_filter._scope_path != "/"):
                            logger.warning(
                                f"[SCOPE] First page yielded 0 in-scope links but "
                                f"{len(self._scope_rejected_buffer)} were scope-rejected "
                                f"— widening to domain"
                            )
                            self._scope_filter.widen_to_domain()
                            self._scope_filter.log_scope()
                            # Re-check buffered scope-rejected links
                            recovered = 0
                            for link in self._scope_rejected_buffer:
                                if link in self._visited_urls or link in self._queued_urls:
                                    continue
                                if self._scope_filter.accept(link):
                                    queue.append((
                                        link,
                                        depth + 1,
                                        page_data.url,
                                        page_data.section_path.copy()
                                    ))
                                    self._queued_urls.add(link)
                                    recovered += 1
                            page_data.internal_links = [
                                link for link in self._scope_rejected_buffer
                                if self._scope_filter.accept(link)
                            ]
                            logger.info(
                                f"[SCOPE] Recovered {recovered} links after "
                                f"scope widening — queue_size={len(queue)}"
                            )
                    
                    # Rate limiting
                    time.sleep(self.config.delay_between_pages)
                else:
                    logger.info(
                        f"[FRONTIER] Page {url[:60]} returned no data "
                        f"(fail/timeout/dup) — queue_size={len(queue)}"
                    )
            
            # Determine final stop reason
            if self._stop_requested:
                stop_reason = "User requested stop"
            elif not queue:
                stop_reason = "Queue exhausted (all reachable pages crawled)"
            
        finally:
            self._close_browser()
        
        # Filter out skipped (empty/cookie) pages from final results
        good_pages = [p for p in self._pages if not getattr(p, '_skipped', False)]
        skipped_count = len(self._pages) - len(good_pages)
        
        # Build stats
        elapsed = time.time() - self._start_time
        stats = {
            'pages_crawled': len(good_pages),
            'pages_skipped': skipped_count,
            'pages_failed': len(self._errors),
            'expandables_clicked': self._expandables_clicked,
            'links_discovered': self._links_discovered,
            'elapsed_time': round(elapsed, 2),
            'pages_per_second': round(len(good_pages) / elapsed, 2) if elapsed > 0 else 0,
            'stop_reason': stop_reason,
            'scope': scope_desc,
        }
        
        logger.info("=" * 60)
        logger.info(f"DEEP CRAWL COMPLETE")
        logger.info(f"Pages crawled: {stats['pages_crawled']}")
        logger.info(f"Pages skipped (empty/cookie): {stats['pages_skipped']}")
        logger.info(f"Pages failed: {stats['pages_failed']}")
        logger.info(f"Expandables clicked: {stats['expandables_clicked']}")
        logger.info(f"Elapsed time: {stats['elapsed_time']}s")
        logger.info(f"Stop reason: {stop_reason}")
        logger.info("=" * 60)
        
        return DeepCrawlResult(
            pages=good_pages,
            stats=stats,
            errors=self._errors.copy(),
            hierarchy=self._build_hierarchy()
        )
    
    def _build_hierarchy(self) -> Dict:
        """Build a tree structure from crawled pages."""
        hierarchy = {'root': {}, 'pages': {}}
        
        for page in self._pages:
            # Add to flat lookup
            hierarchy['pages'][page.url] = {
                'title': page.title,
                'section_path': page.section_path,
                'depth': page.depth,
            }
            
            # Build tree
            if page.section_path:
                current = hierarchy['root']
                for section in page.section_path:
                    if section not in current:
                        current[section] = {'_pages': [], '_children': {}}
                    current[section]['_pages'].append(page.url)
                    current = current[section]['_children']
        
        return hierarchy
    
    def export_json(self, result: DeepCrawlResult, filepath: str) -> str:
        """Export results to JSON."""
        output_path = Path(filepath)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        data = {
            'stats': result.stats,
            'hierarchy': result.hierarchy,
            'pages': [p.to_dict() for p in result.pages],
            'errors': result.errors,
        }
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        
        return str(output_path.absolute())
    
    def export_csv(self, result: DeepCrawlResult, filepath: str) -> str:
        """Export results to CSV."""
        import csv
        
        output_path = Path(filepath)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        if not result.pages:
            return str(output_path.absolute())
        
        rows = [p.to_flat_dict() for p in result.pages]
        
        with open(output_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.DictWriter(f, fieldnames=rows[0].keys())
            writer.writeheader()
            writer.writerows(rows)
        
        return str(output_path.absolute())
    
    def export_docx(self, result: DeepCrawlResult, filepath: str) -> str:
        """
        Export crawl results to a professionally formatted Word document.
        
        Each page includes:
        - Page Title
        - Page URL  
        - Breadcrumb/Section Path
        - Headings (H1-H6)
        - Visible Text Content
        
        Args:
            result: DeepCrawlResult to export
            filepath: Output file path
            
        Returns:
            Absolute path to the created file
        """
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        output_path = Path(filepath)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        doc = Document()
        
        # Configure default style
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(10)
        
        # Cover page / Summary
        title = doc.add_heading('Web Crawl Report', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Summary stats
        summary_items = [
            f"Crawl Scope: {result.stats.get('scope', 'N/A')}",
            f"Total Pages Crawled: {result.stats.get('pages_crawled', 0)}",
            f"Pages Failed: {result.stats.get('pages_failed', 0)}",
            f"Expandables Clicked: {result.stats.get('expandables_clicked', 0)}",
            f"Elapsed Time: {result.stats.get('elapsed_time', 0)}s",
            f"Speed: {result.stats.get('pages_per_second', 0):.2f} pages/sec",
            f"Stop Reason: {result.stats.get('stop_reason', 'N/A')}",
        ]
        
        for item in summary_items:
            p = doc.add_paragraph(item)
            p.paragraph_format.space_after = Pt(2)
        
        doc.add_page_break()
        
        # Per-page content
        for idx, page in enumerate(result.pages):
            # Page title
            heading_text = page.title if page.title else page.url
            doc.add_heading(heading_text[:100], level=1)
            
            # URL
            url_para = doc.add_paragraph()
            url_run = url_para.add_run(page.url)
            url_run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
            url_run.font.size = Pt(9)
            url_para.paragraph_format.space_after = Pt(4)
            
            # Section path / breadcrumb
            if page.section_path:
                path_para = doc.add_paragraph()
                path_label = path_para.add_run('Section: ')
                path_label.bold = True
                path_label.font.size = Pt(9)
                path_value = path_para.add_run(' > '.join(page.section_path))
                path_value.font.size = Pt(9)
                path_value.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            
            if page.breadcrumb:
                bc_para = doc.add_paragraph()
                bc_label = bc_para.add_run('Breadcrumb: ')
                bc_label.bold = True
                bc_label.font.size = Pt(9)
                bc_value = bc_para.add_run(' > '.join(page.breadcrumb))
                bc_value.font.size = Pt(9)
            
            # Headings (H1-H6)
            if page.headings:
                doc.add_heading('Headings', level=2)
                for level_tag, heading_list in page.headings.items():
                    for h_text in heading_list[:10]:  # Limit headings per level
                        p = doc.add_paragraph(style='List Bullet')
                        tag_run = p.add_run(f'[{level_tag.upper()}] ')
                        tag_run.bold = True
                        tag_run.font.size = Pt(9)
                        text_run = p.add_run(h_text[:200])
                        text_run.font.size = Pt(9)
            
            # Main text content
            if page.text_content:
                doc.add_heading('Content', level=2)
                content = page.text_content
                # Limit to 15000 chars per page
                if len(content) > 15000:
                    content = content[:15000] + '\n\n[... content truncated ...]'
                
                # Write in chunks for readability
                chunks = [content[i:i+2000] for i in range(0, len(content), 2000)]
                for chunk in chunks:
                    p = doc.add_paragraph(chunk)
                    p.paragraph_format.space_after = Pt(4)
                    for run in p.runs:
                        run.font.size = Pt(9)
            
            # Page break between pages (except last)
            if idx < len(result.pages) - 1:
                doc.add_page_break()
        
        doc.save(str(output_path))
        logger.info(f"Exported DOCX to {output_path.absolute()}")
        return str(output_path.absolute())


# Convenience function
def deep_crawl_docs(url: str, max_pages: int = 100, max_depth: int = 5) -> DeepCrawlResult:
    """
    Quick function to deep crawl a documentation site.
    
    Args:
        url: Starting URL
        max_pages: Maximum pages to crawl
        max_depth: Maximum depth to crawl
        
    Returns:
        DeepCrawlResult
    """
    config = DeepCrawlConfig(max_pages=max_pages, max_depth=max_depth)
    crawler = DeepDocCrawler(config)
    return crawler.crawl(url)


# =============================================================================
# CLI INTERFACE
# =============================================================================
def main():
    """
    CLI entry point for the deep web crawler.
    
    Usage:
        python -m crawler.deep_crawler <url> [options]
        
    Examples:
        python -m crawler.deep_crawler https://docs.example.com --max_pages 100
        python -m crawler.deep_crawler https://example.com/docs --max_depth 3 --output_json results.json
        python -m crawler.deep_crawler https://site.com --no_js --output_csv data.csv
    """
    import argparse
    import sys
    
    # Configure logging for CLI
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s | %(levelname)-7s | %(message)s',
        datefmt='%H:%M:%S'
    )
    
    parser = argparse.ArgumentParser(
        description='Advanced Web Crawler - CLI-first production crawler',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
Examples:
  python -m crawler.deep_crawler https://docs.example.com
  python -m crawler.deep_crawler https://example.com/docs --max_pages 100 --max_depth 3
  python -m crawler.deep_crawler https://site.com --output_json results.json --output_csv data.csv
  python -m crawler.deep_crawler https://site.com --no_js --rate 0.5 --output_docx report.docx

Scope Rules:
  - https://example.com        -> Crawls entire domain
  - https://example.com/docs   -> Crawls only /docs/** subtree

Safety Limits (enforced strictly):
  - Per-page timeout: 20 seconds (configurable)
  - Per-page max clicks: 50 (to prevent infinite expansion)
  - Default MAX_PAGES: 150
  - Default MAX_DEPTH: 5
        '''
    )
    
    # Required
    parser.add_argument('url', help='Starting URL to crawl')
    
    # Crawl limits
    parser.add_argument('--max_depth', type=int, default=5,
                        help='Maximum crawl depth (default: 5)')
    parser.add_argument('--max_pages', type=int, default=150,
                        help='Maximum pages to crawl (default: 150)')
    parser.add_argument('--timeout', type=int, default=20,
                        help='Per-page timeout in seconds (default: 20)')
    
    # Rate limiting
    parser.add_argument('--rate', type=float, default=1.0,
                        help='Delay between page requests in seconds (default: 1.0)')
    
    # JS handling
    parser.add_argument('--no_js', action='store_true',
                        help='Disable JavaScript rendering (static HTML only)')
    parser.add_argument('--max_clicks', type=int, default=50,
                        help='Max expandable clicks per page (default: 50)')
    
    # Output files
    parser.add_argument('--output_json', type=str, default=None,
                        help='Output JSON file path')
    parser.add_argument('--output_csv', type=str, default=None,
                        help='Output CSV file path')
    parser.add_argument('--output_docx', type=str, default=None,
                        help='Output DOCX file path')
    
    # Other options
    parser.add_argument('--headless', type=bool, default=True,
                        help='Run browser in headless mode (default: True)')
    parser.add_argument('-v', '--verbose', action='store_true',
                        help='Enable verbose/debug logging')
    
    args = parser.parse_args()
    
    # Set debug logging if verbose
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    
    # Validate URL
    if not args.url.startswith(('http://', 'https://')):
        logger.error("URL must start with http:// or https://")
        sys.exit(1)
    
    # Build configuration
    config = DeepCrawlConfig(
        max_pages=args.max_pages,
        max_depth=args.max_depth,
        timeout=args.timeout * 1000,  # Convert to milliseconds
        delay_between_pages=args.rate,
        max_clicks_per_page=args.max_clicks,
        headless=args.headless,
        enable_static_fallback=True,
    )
    
    # If --no_js, we'll use the static-only approach
    use_static_only = args.no_js
    
    logger.info("=" * 60)
    logger.info("ADVANCED WEB CRAWLER - CLI MODE")
    logger.info("=" * 60)
    logger.info(f"URL: {args.url}")
    logger.info(f"Max Pages: {args.max_pages}")
    logger.info(f"Max Depth: {args.max_depth}")
    logger.info(f"Timeout: {args.timeout}s per page")
    logger.info(f"Rate Limit: {args.rate}s between requests")
    logger.info(f"JS Rendering: {'Disabled' if use_static_only else 'Enabled'}")
    logger.info(f"Max Clicks/Page: {args.max_clicks}")
    logger.info("=" * 60)
    
    try:
        if use_static_only:
            # Use the regular WebCrawler with JS disabled
            from .crawler import WebCrawler, CrawlConfig
            
            crawl_config = CrawlConfig(
                max_depth=args.max_depth,
                max_pages=args.max_pages,
                timeout=args.timeout,
                requests_per_second=1.0 / args.rate if args.rate > 0 else 1.0,
                enable_js_rendering=False,
            )
            
            crawler = WebCrawler(crawl_config)
            
            def progress_callback(pages, url, stats):
                logger.info(f"[{pages}/{args.max_pages}] {url[:70]}...")
            
            crawler.set_progress_callback(progress_callback)
            result = crawler.crawl(args.url)
            
            # Export
            if args.output_json:
                crawler.export_json(result, args.output_json)
                logger.info(f"JSON exported to: {args.output_json}")
            
            if args.output_csv:
                crawler.export_csv(result, args.output_csv)
                logger.info(f"CSV exported to: {args.output_csv}")
            
            if args.output_docx:
                crawler.export_docx(result, args.output_docx)
                logger.info(f"DOCX exported to: {args.output_docx}")
            
            # Summary
            logger.info("=" * 60)
            logger.info("CRAWL COMPLETE")
            logger.info(f"Pages crawled: {result.stats.get('pages_crawled', 0)}")
            logger.info(f"Pages failed: {result.stats.get('pages_failed', 0)}")
            logger.info(f"Elapsed time: {result.stats.get('elapsed_time', 0)}s")
            logger.info("=" * 60)
            
        else:
            # Use DeepDocCrawler for JS-rendered sites
            crawler = DeepDocCrawler(config)
            
            def progress_callback(pages, url, stats):
                pass  # Already logged in crawl method
            
            crawler.set_progress_callback(progress_callback)
            result = crawler.crawl(args.url)
            
            # Export
            if args.output_json:
                crawler.export_json(result, args.output_json)
                logger.info(f"JSON exported to: {args.output_json}")
            
            if args.output_csv:
                crawler.export_csv(result, args.output_csv)
                logger.info(f"CSV exported to: {args.output_csv}")
            
            if args.output_docx:
                crawler.export_docx(result, args.output_docx)
                logger.info(f"DOCX exported to: {args.output_docx}")
        
        # Default export if none specified
        if not args.output_json and not args.output_csv and not args.output_docx:
            default_json = 'crawl_output.json'
            if use_static_only:
                crawler.export_json(result, default_json)
            else:
                crawler.export_json(result, default_json)
            logger.info(f"No output specified, exported to: {default_json}")
        
    except KeyboardInterrupt:
        logger.info("\nCrawl interrupted by user (Ctrl+C)")
        sys.exit(0)
    except Exception as e:
        logger.error(f"Crawl failed: {e}")
        if args.verbose:
            import traceback
            traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
