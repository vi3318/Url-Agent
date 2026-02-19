"""
Web Crawler
Main crawler implementation with BFS/DFS traversal, JS rendering fallback,
and comprehensive crawling logic.
"""

import logging
import time
import json
import csv
import platform
from typing import List, Dict, Set, Optional, Callable, Generator
from dataclasses import dataclass, field
from collections import deque
from urllib.parse import urlparse
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
import threading

import requests
from playwright.sync_api import sync_playwright, TimeoutError as PlaywrightTimeout

from .robots import RobotsHandler
from .scraper import PageScraper, PageData, detect_js_required
from .scope_filter import ScopeFilter
from .utils import (
    URLNormalizer, RateLimiter, RetryHandler,
    ProgressTracker, is_valid_url, extract_domain
)

logger = logging.getLogger(__name__)


def _detect_platform() -> str:
    """Return the sec-ch-ua-platform value for the current OS."""
    system = platform.system()
    if system == "Darwin":
        return "macOS"
    elif system == "Windows":
        return "Windows"
    else:
        return "Linux"


@dataclass
class CrawlConfig:
    """
    Configuration for the web crawler.
    """
    # Crawl limits
    max_depth: int = 3
    max_pages: int = 1000
    timeout: int = 30
    
    # Rate limiting
    requests_per_second: float = 1.0
    respect_robots: bool = True
    
    # User agent – realistic Chrome UA so sites don't reject us
    user_agent: str = (
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
        "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36"
    )
    
    # Retry settings
    max_retries: int = 3
    retry_delay: float = 1.0
    
    # JavaScript rendering
    enable_js_rendering: bool = True
    js_render_timeout: int = 45000  # milliseconds (45s for slow enterprise sites)
    auto_detect_js: bool = True
    
    # Content extraction
    extract_images: bool = False
    extract_external_links: bool = True
    max_text_length: int = None
    
    # Threading
    max_workers: int = 1  # Keep at 1 for politeness, increase cautiously
    
    # Output
    output_dir: str = "./output"

    # Scope-filter params (populated by CrawlerRunConfig converters)
    deny_patterns: List[str] = field(default_factory=list)
    strip_all_queries: bool = False
    
    def to_dict(self) -> dict:
        """Convert config to dictionary."""
        return {
            'max_depth': self.max_depth,
            'max_pages': self.max_pages,
            'timeout': self.timeout,
            'requests_per_second': self.requests_per_second,
            'respect_robots': self.respect_robots,
            'user_agent': self.user_agent,
            'max_retries': self.max_retries,
            'enable_js_rendering': self.enable_js_rendering,
            'auto_detect_js': self.auto_detect_js,
            'max_workers': self.max_workers
        }


@dataclass
class CrawlResult:
    """
    Result of a crawl operation.
    """
    pages: List[PageData] = field(default_factory=list)
    stats: Dict = field(default_factory=dict)
    errors: List[Dict] = field(default_factory=list)
    config: Dict = field(default_factory=dict)
    scope_info: Dict = field(default_factory=dict)


class WebCrawler:
    """
    Production-grade web crawler with support for static and JavaScript-rendered pages.
    """
    
    def __init__(self, config: CrawlConfig = None):
        """
        Initialize the web crawler.
        
        Args:
            config: Crawler configuration
        """
        self.config = config or CrawlConfig()
        
        # Initialize components
        self.robots_handler = RobotsHandler(
            user_agent=self.config.user_agent,
            respect_robots=self.config.respect_robots
        )
        
        self.url_normalizer = URLNormalizer()
        
        self.rate_limiter = RateLimiter(
            requests_per_second=self.config.requests_per_second
        )
        
        self.retry_handler = RetryHandler(
            max_retries=self.config.max_retries,
            base_delay=self.config.retry_delay
        )
        
        self.progress = ProgressTracker()
        
        # Session for HTTP requests
        self.session = self._create_session()
        
        # Playwright browser (lazy initialization)
        self._playwright = None
        self._browser = None
        
        # Thread safety
        self._lock = threading.Lock()
        
        # Scope filter (initialized per-crawl in .crawl())
        self._scope_filter: Optional[ScopeFilter] = None
        
        # Crawl state
        self._visited_urls: Set[str] = set()
        self._queued_urls: Set[str] = set()   # frontier dedup — prevents double-enqueue
        self._failed_urls: Set[str] = set()
        self._pages: List[PageData] = []
        self._errors: List[Dict] = []
        
        # Callback for progress updates
        self._progress_callback: Optional[Callable] = None
        
        # Stop flag
        self._stop_requested = False
    
    def _create_session(self) -> requests.Session:
        """Create configured requests session with realistic browser headers."""
        session = requests.Session()
        session.headers.update({
            'User-Agent': self.config.user_agent,
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.9',
            'Accept-Encoding': 'gzip, deflate',
            'Connection': 'keep-alive',
            'Upgrade-Insecure-Requests': '1',
            'Sec-Fetch-Dest': 'document',
            'Sec-Fetch-Mode': 'navigate',
            'Sec-Fetch-Site': 'none',
            'Sec-Fetch-User': '?1',
            'Cache-Control': 'max-age=0',
            'sec-ch-ua': '"Chromium";v="122", "Not(A:Brand";v="24", "Google Chrome";v="122"',
            'sec-ch-ua-mobile': '?0',
            'sec-ch-ua-platform': f'"{_detect_platform()}"',
        })
        return session
    
    def _init_playwright(self) -> None:
        """Initialize Playwright browser for JS rendering."""
        if self._playwright is None and self.config.enable_js_rendering:
            try:
                self._playwright = sync_playwright().start()
                self._browser = self._playwright.chromium.launch(
                    headless=True,
                    args=['--disable-gpu', '--no-sandbox', '--disable-dev-shm-usage']
                )
                logger.info("Playwright browser initialized for JS rendering")
            except Exception as e:
                logger.warning(f"Failed to initialize Playwright: {e}")
                self.config.enable_js_rendering = False
    
    def _close_playwright(self) -> None:
        """Close Playwright browser."""
        if self._browser:
            try:
                self._browser.close()
            except Exception:
                pass
            self._browser = None
        
        if self._playwright:
            try:
                self._playwright.stop()
            except Exception:
                pass
            self._playwright = None
    
    def set_progress_callback(self, callback: Callable) -> None:
        """
        Set callback for progress updates.
        
        Args:
            callback: Function(pages_crawled, current_url, stats)
        """
        self._progress_callback = callback
    
    def stop(self) -> None:
        """Request the crawler to stop."""
        self._stop_requested = True
        logger.info("Stop requested")
    
    def _fetch_page_static(self, url: str) -> tuple:
        """
        Fetch page using requests (static rendering).
        
        Returns:
            Tuple of (html, status_code, content_type, error)
        """
        try:
            response = self.session.get(
                url,
                timeout=self.config.timeout,
                allow_redirects=True
            )
            
            content_type = response.headers.get('Content-Type', '')
            
            # Only process HTML content (accept text/html, application/xhtml+xml, or missing content-type)
            ct_lower = content_type.lower()
            is_html = (
                'text/html' in ct_lower or
                'xhtml' in ct_lower or
                not content_type  # No content-type header — try anyway
            )
            if not is_html:
                return None, response.status_code, content_type, "Not HTML content"
            
            return response.text, response.status_code, content_type, None
            
        except requests.Timeout:
            return None, 0, "", "Request timeout"
        except requests.RequestException as e:
            return None, 0, "", str(e)
    
    def _fetch_page_js(self, url: str) -> tuple:
        """
        Fetch page using Playwright (JavaScript rendering).
        
        Returns:
            Tuple of (html, status_code, content_type, error)
        """
        if not self.config.enable_js_rendering or not self._browser:
            return None, 0, "", "JS rendering not available"
        
        try:
            context = self._browser.new_context(
                user_agent=self.config.user_agent,
                viewport={'width': 1920, 'height': 1080},
                locale='en-US',
                timezone_id='America/New_York',
                java_script_enabled=True,
            )
            page = context.new_page()
            
            try:
                # Use 'load' instead of 'networkidle' to avoid hanging on
                # sites with constant background requests (analytics, chat, etc.)
                response = page.goto(
                    url,
                    timeout=self.config.js_render_timeout,
                    wait_until='load'
                )
                
                if response is None:
                    return None, 0, "", "No response from page"
                
                # Wait for DOM to be ready
                page.wait_for_load_state('domcontentloaded')
                
                # Give JS frameworks time to render (React, Vue, Angular, etc.)
                try:
                    page.wait_for_timeout(3000)
                except Exception:
                    pass
                
                # Try to wait for network to settle, but with a short timeout
                try:
                    page.wait_for_load_state('networkidle', timeout=5000)
                except Exception:
                    # Network didn't settle in 5s, proceed anyway
                    pass
                
                # Scroll down to trigger lazy loading
                try:
                    page.evaluate("window.scrollTo(0, document.body.scrollHeight / 2)")
                    page.wait_for_timeout(1000)
                    page.evaluate("window.scrollTo(0, document.body.scrollHeight)")
                    page.wait_for_timeout(1000)
                except Exception:
                    pass
                
                html = page.content()
                status = response.status
                content_type = response.headers.get('content-type', 'text/html')
                
                return html, status, content_type, None
                
            finally:
                page.close()
                context.close()
                
        except PlaywrightTimeout:
            return None, 0, "", "JS rendering timeout"
        except Exception as e:
            return None, 0, "", f"JS rendering error: {str(e)}"
    
    def _fetch_page(self, url: str, use_js: bool = False) -> tuple:
        """
        Fetch a page with optional JS rendering fallback.
        
        Args:
            url: URL to fetch
            use_js: Force JavaScript rendering
            
        Returns:
            Tuple of (html, status_code, content_type, used_js, error)
        """
        # Rate limiting
        self.rate_limiter.wait(url)
        
        # Check robots.txt crawl delay
        crawl_delay = self.robots_handler.get_crawl_delay(url)
        if crawl_delay:
            self.rate_limiter.set_crawl_delay(url, crawl_delay)
        
        # First try static fetch
        if not use_js:
            html, status, content_type, error = self._fetch_page_static(url)
            
            if html and self.config.auto_detect_js and detect_js_required(html):
                logger.debug(f"JS rendering detected as needed for {url}")
                use_js = True
            elif html:
                return html, status, content_type, False, None
            elif error:
                # Try JS rendering as fallback
                if self.config.enable_js_rendering:
                    use_js = True
                else:
                    return None, status, content_type, False, error
        
        # JavaScript rendering
        if use_js:
            self._init_playwright()
            html, status, content_type, error = self._fetch_page_js(url)
            return html, status, content_type, True, error
        
        return None, 0, "", False, "Unknown error"
    
    def _process_url(
        self,
        url: str,
        depth: int,
        base_url: str
    ) -> Optional[PageData]:
        """
        Process a single URL: fetch and scrape.
        
        Args:
            url: URL to process
            depth: Current crawl depth
            base_url: Original crawl root URL
            
        Returns:
            PageData or None if failed
        """
        # Normalize URL
        normalized_url = self.url_normalizer.normalize(url, base_url)
        if not normalized_url:
            return None
        
        # Check if already visited
        with self._lock:
            if normalized_url in self._visited_urls:
                return None
        
        # Check robots.txt
        if not self.robots_handler.can_fetch(normalized_url):
            logger.debug(f"Blocked by robots.txt: {normalized_url}")
            self.progress.increment_skipped()
            return None
        
        # Scope check — belt-and-suspenders guard (primary filtering is at enqueue time)
        if self._scope_filter and not self._scope_filter.accept(normalized_url):
            logger.debug(f"[SCOPE] Rejected at process-url gate: {normalized_url}")
            self.progress.increment_skipped()
            return None
        
        # Mark visited AFTER scope + robots checks so rejected URLs aren't permanently eaten
        with self._lock:
            self._visited_urls.add(normalized_url)
        
        logger.info(f"Crawling [{depth}]: {normalized_url}")
        
        # Fetch page with retry
        try:
            html, status, content_type, used_js, error = self.retry_handler.execute_with_retry(
                self._fetch_page,
                normalized_url
            )
        except Exception as e:
            error = str(e)
            html = None
            status = 0
            content_type = ""
            used_js = False
        
        if error or not html:
            self.progress.increment_failed()
            self._errors.append({
                'url': normalized_url,
                'error': error or "Empty response",
                'depth': depth
            })
            return None
        
        # Scrape content
        scraper = PageScraper(
            base_url=base_url,
            extract_images=self.config.extract_images,
            extract_external_links=self.config.extract_external_links,
            max_text_length=self.config.max_text_length
        )
        
        page_data = scraper.scrape(
            html=html,
            url=normalized_url,
            status_code=status,
            content_type=content_type,
            depth=depth
        )
        
        self.progress.increment_crawled()
        
        # Progress callback
        if self._progress_callback:
            try:
                self._progress_callback(
                    self.progress.pages_crawled,
                    normalized_url,
                    self.progress.get_stats()
                )
            except Exception:
                pass
        
        return page_data
    
    def crawl(
        self,
        start_url: str,
        strategy: str = "bfs"
    ) -> CrawlResult:
        """
        Crawl a website starting from the given URL.
        
        Args:
            start_url: Starting URL for the crawl
            strategy: Crawling strategy ('bfs' or 'dfs')
            
        Returns:
            CrawlResult with all scraped data
        """
        # Validate URL
        if not is_valid_url(start_url):
            raise ValueError(f"Invalid URL: {start_url}")
        
        # Normalize start URL
        start_url = self.url_normalizer.normalize(start_url) or start_url
        
        # Initialize scope filter for this crawl
        self._scope_filter = ScopeFilter(
            root_url=start_url,
            deny_patterns=self.config.deny_patterns,
            strip_all_queries=self.config.strip_all_queries,
        )
        self._scope_filter.log_scope()
        
        # Get and log scope information
        scope_info = {
            'scope_description': self._scope_filter.scope_description,
            'base_domain': self._scope_filter._root_canon.host if self._scope_filter._root_canon else '',
            'base_path': self._scope_filter._scope_path,
        }
        
        logger.info(f"Starting crawl of {start_url} with {strategy.upper()} strategy")
        logger.info(f"Crawl scope: {scope_info['scope_description']}")
        logger.info(f"Config: max_depth={self.config.max_depth}, max_pages={self.config.max_pages}")
        
        # Reset state
        self._visited_urls.clear()
        self._queued_urls.clear()
        self._failed_urls.clear()
        self._pages.clear()
        self._errors.clear()
        self._stop_requested = False
        
        # Start progress tracking
        self.progress = ProgressTracker()
        self.progress.start()
        
        try:
            if strategy.lower() == "bfs":
                self._crawl_bfs(start_url)
            elif strategy.lower() == "dfs":
                self._crawl_dfs(start_url)
            else:
                raise ValueError(f"Unknown strategy: {strategy}")
        finally:
            self.progress.finish()
            self._close_playwright()
        
        # Build result
        result = CrawlResult(
            pages=self._pages.copy(),
            stats=self.progress.get_stats(),
            errors=self._errors.copy(),
            config=self.config.to_dict(),
            scope_info=scope_info
        )
        
        logger.info(f"Crawl complete. Stats: {result.stats}")
        
        return result
    
    def _crawl_bfs(self, start_url: str) -> None:
        """Breadth-first crawl implementation."""
        # Queue: (url, depth)
        queue = deque([(start_url, 0)])
        self._queued_urls.add(start_url)
        
        while queue and not self._stop_requested:
            # Check page limit
            if self.progress.pages_crawled >= self.config.max_pages:
                logger.info(f"Reached max pages limit: {self.config.max_pages}")
                break
            
            url, depth = queue.popleft()
            
            # Check depth limit
            if depth > self.config.max_depth:
                continue
            
            logger.info(f"[BFS] Depth:{depth} | Queue:{len(queue)} | {url[:80]}")
            
            # Process URL
            page_data = self._process_url(url, depth, start_url)
            
            if page_data:
                self._pages.append(page_data)
                
                # Add internal links to queue — scope-filtered at enqueue time
                if depth < self.config.max_depth:
                    new_enqueued = 0
                    rejected_scope = 0
                    for link in page_data.internal_links:
                        normalized = self.url_normalizer.normalize(link, start_url)
                        if not normalized:
                            continue
                        if normalized in self._visited_urls or normalized in self._queued_urls:
                            continue
                        # Scope gate — reject before it ever enters the queue
                        if self._scope_filter and not self._scope_filter.accept(normalized):
                            rejected_scope += 1
                            continue
                        queue.append((normalized, depth + 1))
                        self._queued_urls.add(normalized)
                        new_enqueued += 1
                    
                    logger.info(
                        f"[FRONTIER] {url[:60]} → "
                        f"links={len(page_data.internal_links)} "
                        f"enqueued={new_enqueued} "
                        f"scope_rejected={rejected_scope} "
                        f"queue_size={len(queue)}"
                    )
            else:
                logger.info(
                    f"[FRONTIER] {url[:60]} returned no data — queue_size={len(queue)}"
                )
    
    def _crawl_dfs(self, start_url: str) -> None:
        """Depth-first crawl implementation."""
        # Stack: (url, depth)
        stack = [(start_url, 0)]
        self._queued_urls.add(start_url)
        
        while stack and not self._stop_requested:
            # Check page limit
            if self.progress.pages_crawled >= self.config.max_pages:
                logger.info(f"Reached max pages limit: {self.config.max_pages}")
                break
            
            url, depth = stack.pop()
            
            # Check depth limit
            if depth > self.config.max_depth:
                continue
            
            # Process URL
            page_data = self._process_url(url, depth, start_url)
            
            if page_data:
                self._pages.append(page_data)
                
                # Add internal links to stack — scope-filtered at enqueue time
                if depth < self.config.max_depth:
                    for link in reversed(page_data.internal_links):
                        normalized = self.url_normalizer.normalize(link, start_url)
                        if not normalized:
                            continue
                        if normalized in self._visited_urls or normalized in self._queued_urls:
                            continue
                        if self._scope_filter and not self._scope_filter.accept(normalized):
                            continue
                        stack.append((normalized, depth + 1))
                        self._queued_urls.add(normalized)
    
    def export_json(self, result: CrawlResult, filepath: str) -> str:
        """
        Export crawl results to JSON.
        
        Args:
            result: Crawl result to export
            filepath: Output file path
            
        Returns:
            Absolute path to the created file
        """
        output_path = Path(filepath)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        data = {
            'metadata': {
                'total_pages': len(result.pages),
                'crawl_stats': result.stats,
                'config': result.config,
                'errors_count': len(result.errors)
            },
            'pages': [page.to_dict() for page in result.pages],
            'errors': result.errors
        }
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        
        logger.info(f"Exported JSON to {output_path.absolute()}")
        return str(output_path.absolute())
    
    def export_csv(self, result: CrawlResult, filepath: str) -> str:
        """
        Export crawl results to CSV.
        
        Args:
            result: Crawl result to export
            filepath: Output file path
            
        Returns:
            Absolute path to the created file
        """
        output_path = Path(filepath)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        if not result.pages:
            logger.warning("No pages to export")
            return str(output_path.absolute())
        
        # Get field names from first page
        fieldnames = list(result.pages[0].to_flat_dict().keys())
        
        with open(output_path, 'w', encoding='utf-8', newline='') as f:
            writer = csv.DictWriter(f, fieldnames=fieldnames, extrasaction='ignore')
            writer.writeheader()
            
            for page in result.pages:
                writer.writerow(page.to_flat_dict())
        
        logger.info(f"Exported CSV to {output_path.absolute()}")
        return str(output_path.absolute())
    
    def export_docx(self, result: CrawlResult, filepath: str) -> str:
        """
        Export crawl results to a professionally formatted Word document.
        
        Args:
            result: Crawl result to export
            filepath: Output file path
            
        Returns:
            Absolute path to the created file
        """
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.section import WD_ORIENT
        
        output_path = Path(filepath)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        doc = Document()
        
        # -- Styles ----------------------------------------------------------
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(10)
        style.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        
        # -- Cover / Summary -------------------------------------------------
        title = doc.add_heading('Web Crawl Report', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        scope_desc = result.scope_info.get('scope_description', 'N/A') if result.scope_info else 'N/A'
        summary_items = [
            f"Total Pages Crawled: {len(result.pages)}",
            f"Crawl Scope: {scope_desc}",
            f"Pages Failed: {result.stats.get('pages_failed', 0)}",
            f"Elapsed Time: {result.stats.get('elapsed_time', 0):.1f}s",
            f"Speed: {result.stats.get('pages_per_second', 0):.2f} pages/sec",
        ]
        for item in summary_items:
            p = doc.add_paragraph(item)
            p.paragraph_format.space_after = Pt(2)
        
        doc.add_page_break()
        
        # -- Per-page content -------------------------------------------------
        for idx, page in enumerate(result.pages):
            # Page title
            heading_text = page.title if page.title else page.url
            doc.add_heading(heading_text, level=1)
            
            # URL
            url_para = doc.add_paragraph()
            url_run = url_para.add_run(page.url)
            url_run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
            url_run.font.size = Pt(9)
            url_para.paragraph_format.space_after = Pt(4)
            
            # Meta description
            if page.meta_description:
                meta_para = doc.add_paragraph()
                meta_label = meta_para.add_run('Meta Description: ')
                meta_label.bold = True
                meta_label.font.size = Pt(9)
                meta_value = meta_para.add_run(page.meta_description)
                meta_value.font.size = Pt(9)
                meta_value.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            
            # Headings (H1–H6)
            if page.headings:
                doc.add_heading('Headings', level=2)
                for level_tag, heading_list in page.headings.items():
                    for h_text in heading_list:
                        p = doc.add_paragraph(style='List Bullet')
                        tag_run = p.add_run(f'[{level_tag.upper()}] ')
                        tag_run.bold = True
                        tag_run.font.size = Pt(9)
                        text_run = p.add_run(h_text)
                        text_run.font.size = Pt(9)
            
            # Main text content
            if page.text_content:
                doc.add_heading('Content', level=2)
                # Split into paragraphs to avoid a single massive block
                # and to keep the document readable
                content = page.text_content
                # Limit to 15 000 chars per page to prevent giant docs
                if len(content) > 15000:
                    content = content[:15000] + '\n\n[... content truncated ...]'
                
                # Write in chunks of ~2000 chars to avoid single huge paragraphs
                chunks = [content[i:i+2000] for i in range(0, len(content), 2000)]
                for chunk in chunks:
                    p = doc.add_paragraph(chunk)
                    p.paragraph_format.space_after = Pt(4)
                    for run in p.runs:
                        run.font.size = Pt(9)
            
            # Separator between pages (page break except after last)
            if idx < len(result.pages) - 1:
                doc.add_page_break()
        
        doc.save(str(output_path))
        logger.info(f"Exported DOCX to {output_path.absolute()}")
        return str(output_path.absolute())
    
    def __enter__(self):
        """Context manager entry."""
        return self
    
    def __exit__(self, exc_type, exc_val, exc_tb):
        """Context manager exit."""
        self._close_playwright()
        if self.session:
            self.session.close()


def crawl_website(
    url: str,
    max_depth: int = 3,
    max_pages: int = 100,
    output_json: str = None,
    output_csv: str = None,
    progress_callback: Callable = None,
    **kwargs
) -> CrawlResult:
    """
    Convenience function to crawl a website.
    
    Args:
        url: Starting URL
        max_depth: Maximum crawl depth
        max_pages: Maximum pages to crawl
        output_json: JSON output file path
        output_csv: CSV output file path
        progress_callback: Progress update callback
        **kwargs: Additional CrawlConfig parameters
        
    Returns:
        CrawlResult with scraped data
    """
    config = CrawlConfig(
        max_depth=max_depth,
        max_pages=max_pages,
        **kwargs
    )
    
    with WebCrawler(config) as crawler:
        if progress_callback:
            crawler.set_progress_callback(progress_callback)
        
        result = crawler.crawl(url)
        
        if output_json:
            crawler.export_json(result, output_json)
        
        if output_csv:
            crawler.export_csv(result, output_csv)
        
        return result


# CLI interface
if __name__ == "__main__":
    import argparse
    
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )
    
    parser = argparse.ArgumentParser(description='Web Crawler')
    parser.add_argument('url', help='Starting URL to crawl')
    parser.add_argument('--depth', type=int, default=3, help='Maximum crawl depth')
    parser.add_argument('--pages', type=int, default=100, help='Maximum pages to crawl')
    parser.add_argument('--output-json', default='output.json', help='JSON output file')
    parser.add_argument('--output-csv', default='output.csv', help='CSV output file')
    parser.add_argument('--rate', type=float, default=1.0, help='Requests per second')
    parser.add_argument('--no-js', action='store_true', help='Disable JavaScript rendering')
    
    args = parser.parse_args()
    
    def progress(pages, url, stats):
        print(f"[{pages}] {url}")
    
    result = crawl_website(
        url=args.url,
        max_depth=args.depth,
        max_pages=args.pages,
        output_json=args.output_json,
        output_csv=args.output_csv,
        requests_per_second=args.rate,
        enable_js_rendering=not args.no_js,
        progress_callback=progress
    )
    
    print(f"\nCrawl complete!")
    print(f"Pages crawled: {result.stats['pages_crawled']}")
    print(f"Errors: {len(result.errors)}")
